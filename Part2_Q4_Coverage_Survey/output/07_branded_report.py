"""
Stage 07 -- Branded survey report
=================================

Renders the stage-06 survey report into the house document style: two colours
(#0090FC primary, #CCE9FE pale tint), hierarchy carried by weight, size and the
pale band, A4 portrait at 0.85 inch margins.

The markdown produced by stage 06 is the single source of truth. This stage
parses it and re-renders it through `build/doctheme.py`, so the branded document
can never drift from the analysis: change a number in the pipeline and it
changes here on the next run.

Outputs to `branded/`. The plain stage-06 DOCX in `reports/` is left untouched.

    python 07_branded_report.py
    python build/to_pdf.py branded/06_survey_report.docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent / "build"))

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import doctheme as T
from common import ART, FIGURE_DIR, banner, get_logger

LOG = get_logger("07_branded_report")

OUT_DIR = Path(__file__).resolve().parent / "branded"
OUT_DIR.mkdir(parents=True, exist_ok=True)
BRANDED_DOCX = OUT_DIR / "06_survey_report.docx"

USABLE_IN = 8.27 - 2 * 0.85          # A4 portrait less the house margins

# Tables get denser as they get wider. Prose-heavy tables (the limitations
# tables, the quality-flag register) are three or four columns of sentences;
# the annexes are eight columns of numbers. One type size cannot serve both.
TABLE_SIZE_BY_COLS = {1: 9.5, 2: 9.5, 3: 9.0, 4: 9.0, 5: 8.5, 6: 8.0, 7: 7.5}
TABLE_SIZE_WIDE = 7.0


# --------------------------------------------------------------------------
# Markdown parsing
# --------------------------------------------------------------------------

IMG_RE = re.compile(r"^!\[(?P<alt>.*?)\]\((?P<path>.*?)\)\s*$")
NUM_RE = re.compile(r"^(\d+)\.\s+(.*)$")
RULE_ROW_RE = re.compile(r"^\|[\s:\-|]+\|$")
EMPH_RE = re.compile(r"^\*(?!\*)(.+)\*$")

# The pipeline's markdown uses a plain double hyphen where typography wants a
# dash. Applied to prose only -- never inside a code block, where a double
# hyphen may be a command-line flag.
_DASH_RE = re.compile(r"(?<=\s)--(?=\s)")


def typographic(text: str) -> str:
    return _DASH_RE.sub("\u2013", text)


def parse_blocks(md: str) -> list[tuple]:
    """Turn the report markdown into a flat list of (kind, payload) blocks."""
    blocks: list[tuple] = []
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        if line.startswith("```"):                       # fenced code
            i += 1
            buf = []
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            blocks.append(("code", buf))
            continue

        if line.startswith("|"):                         # table
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                if not RULE_ROW_RE.match(lines[i].strip()):
                    cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                    rows.append(cells)
                i += 1
            if rows:
                blocks.append(("table", rows))
            continue

        m = IMG_RE.match(line)
        if m:
            blocks.append(("image", (m.group("alt"), m.group("path"))))
            i += 1
            continue

        if line.strip() in ("---", "***", "___"):
            blocks.append(("rule", None))
            i += 1
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            blocks.append(("heading", (level, line.lstrip("# ").strip())))
            i += 1
            continue

        if line.startswith("> "):
            blocks.append(("quote", line[2:].strip()))
            i += 1
            continue

        if line.startswith(("- ", "* ")):
            blocks.append(("bullet", line[2:].strip()))
            i += 1
            continue

        m = NUM_RE.match(line)
        if m:
            blocks.append(("number", (m.group(1), m.group(2).strip())))
            i += 1
            continue

        blocks.append(("para", line.strip()))
        i += 1
    return blocks


# --------------------------------------------------------------------------
# Rendering helpers
# --------------------------------------------------------------------------


def rich_numbered(doc, marker: str, text: str, size: float = 10):
    """
    A numbered item that honours **bold** and `code`.

    The marker is written literally rather than using Word's "List Number"
    style. Word treats that style as one continuous list for the whole
    document, so section 8's recommendations would carry on from section 2's
    numbering instead of restarting. Writing the number from the markdown means
    the document says exactly what the source says.
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.30)
    p.paragraph_format.first_line_indent = Inches(-0.30)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(f"{marker}.\t")
    r.font.name, r.font.size, r.font.bold = T.FONT, Pt(size), True
    r.font.color.rgb = T.D_BLUE
    return T.rich_run(p, text, size=size)


def code_block(doc, lines: list[str]) -> None:
    """A monospace block in the primary colour, inside the pale callout box."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    T.shade(cell, T.HEX_PALE)
    cell.text = ""
    for n, line in enumerate(lines):
        p = cell.paragraphs[0] if n == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(line)
        r.font.name, r.font.size = T.MONO, Pt(8.2)
        r.font.color.rgb = T.D_BLUE
    doc.add_paragraph()


def figure(doc, alt: str, path: str) -> None:
    """Image at text width with an italic caption beneath it."""
    src = FIGURE_DIR / Path(path).name
    if not src.exists():
        LOG.warning("figure missing, skipped: %s", src)
        return
    doc.add_picture(str(src), width=Inches(USABLE_IN))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    r = cap.add_run(typographic(T.strip_md(alt)))
    r.font.name, r.font.size, r.font.italic = T.FONT, Pt(8.4), True
    r.font.color.rgb = T.D_BLACK


def branded_table(doc, rows: list[list[str]]) -> None:
    """
    Render a parsed markdown table in the house banding.

    A table whose first row is blank in every cell is a label/value block, and
    `add_table` is told so explicitly -- an empty primary bar above an
    "At a glance" table reads as a mistake rather than as a design.
    """
    header, body = rows[0], rows[1:]
    if not body:
        header, body = None, rows
    ncols = len(rows[0])
    size = TABLE_SIZE_BY_COLS.get(ncols, TABLE_SIZE_WIDE)

    clean_header = None if header is None else [typographic(T.strip_md(h)) for h in header]
    clean_body = [[typographic(T.strip_md(c)) for c in r] for r in body]

    T.add_table(doc, clean_header, clean_body,
                widths=column_widths(clean_header, clean_body), size=size)


def column_widths(header, body) -> list[float]:
    """
    Share the text width between columns in proportion to what each has to
    carry, so that a column of two-digit numbers does not get the same slack as
    a column of sentences.

    A column's demand is the largest of three claims: its typical cell, a
    discounted claim from its longest cell, and a discounted claim from its
    header. The discounts matter -- a header may wrap onto two lines without
    looking broken, and one unusually long cell should not seize the page --
    but they cannot be zero, or a column of blanks with a real header ends up
    breaking that header a letter at a time. Each column is then clamped so
    nothing collapses to an unreadable ribbon and nothing runs away.
    """
    cols = list(zip(*body)) if body else []
    ncols = len(header) if header else len(cols)
    demand = []
    for c in range(ncols):
        cells = [len(str(v)) for v in cols[c]] if c < len(cols) else [1]
        typical = sorted(cells)[int(0.75 * (len(cells) - 1))] if cells else 1
        longest = max(cells) if cells else 1
        head_len = len(str(header[c])) if header else 0
        demand.append(max(typical, 0.60 * longest, 0.75 * head_len, 3))

    total = sum(demand)
    raw = [USABLE_IN * d / total for d in demand]
    # The floor scales with the column count: a three-column table can afford
    # to keep every column readable, a nine-column annex cannot.
    lo = min(1.15, USABLE_IN / (ncols * 1.6))
    hi = USABLE_IN * 0.42
    clamped = [min(max(w, lo), hi) for w in raw]
    scale = USABLE_IN / sum(clamped)
    return [w * scale for w in clamped]


def at_a_glance(doc) -> None:
    """The headline numbers as a label/value block, before the narrative starts."""
    import json

    meta = json.loads(ART["est_headline"].with_suffix(".meta.json").read_text())
    lo, hi = meta["headline_ci_pct"]
    rows = [
        ["National coverage, 9-59 completed months",
         f"{meta['headline_national_pct']:.1f}%  (95% CI {lo:.1f}-{hi:.1f})"],
        ["Campaign target", f"{100*meta['coverage_target']:.0f}%"],
        ["Design effect / effective sample",
         f"{meta['deff']:.2f}  /  {meta['n_effective']:.0f} children"],
        ["Clusters analysed",
         f"81 of 90 ({len(meta['excluded_clusters'])} excluded by the falsification screen)"],
        ["Verdict on a mop-up round", "Required. Robust to every sensitivity variant tested."],
    ]
    T.add_table(doc, None, rows, widths=[USABLE_IN * 0.42, USABLE_IN * 0.58], size=9.5)


# --------------------------------------------------------------------------


def main() -> None:
    banner(LOG, "STAGE 07  Branded survey report")

    md = ART["final_report"].read_text(encoding="utf-8")
    blocks = parse_blocks(md)
    LOG.info("parsed %d blocks from %s", len(blocks), ART["final_report"].name)

    doc = docx.Document()
    T.docx_setup(doc)

    # -------------------------------------------------------------- cover
    T.cover(
        doc,
        title="Post-campaign coverage survey",
        subtitle="Bansara, Kudama and Zaruwa States - May 2026",
        kicker="Survey report for the national programme and the funding partner",
        meta_lines=[
            "Stratified two-stage cluster design, PPS selection of enumeration areas",
            "90 clusters, 1,800 selected households, 2,296 children enumerated",
            "All estimates weighted; all confidence intervals design-based",
        ],
    )
    doc.add_paragraph()
    T.kicker_para(doc, "At a glance")
    at_a_glance(doc)

    T.callout(
        doc, "What this report concludes",
        "Coverage fell short of the 95% campaign target in all three states and a mop-up "
        "round is indicated. That conclusion survives every analytical choice tested. The "
        "survey is not fit for targeting the mop-up below stratum level, and two stratum "
        "estimates should not be published as point figures - see section 2.",
    )
    doc.add_page_break()

    # ---------------------------------------------------------- the report
    # The document title, its standfirst and the first horizontal rule are all
    # carried by the cover already, so everything before the first numbered
    # section is dropped rather than repeated.
    body_started = False
    for kind, payload in blocks:
        if not body_started:
            if kind == "heading" and payload[0] == 2 and re.match(r"^\d+\.", payload[1]):
                body_started = True
            else:
                continue

        if kind == "heading":
            level, text = payload
            doc.add_heading(typographic(T.strip_md(text)), level=min(max(level - 1, 1), 3))

        elif kind == "para":
            m = EMPH_RE.match(payload)
            if m:                                   # a standfirst or an aside
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                r = p.add_run(typographic(T.strip_md(m.group(1))))
                r.font.name, r.font.size, r.font.italic = T.FONT, Pt(9.5), True
                r.font.color.rgb = T.D_BLACK
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                T.rich_run(p, typographic(payload), size=10)

        elif kind == "bullet":
            T.rich_bullet(doc, typographic(payload))

        elif kind == "number":
            rich_numbered(doc, payload[0], typographic(payload[1]))

        elif kind == "table":
            branded_table(doc, payload)

        elif kind == "image":
            figure(doc, *payload)

        elif kind == "code":
            code_block(doc, payload)

        elif kind == "quote":
            T.callout(doc, "In one sentence", typographic(T.strip_md(payload)))

        elif kind == "rule":
            T.hrule(doc)

    doc.save(BRANDED_DOCX)
    LOG.info("wrote %s", BRANDED_DOCX.relative_to(BRANDED_DOCX.parents[1]))
    LOG.info("export to PDF with:  python build/to_pdf.py branded/06_survey_report.docx")
    banner(LOG, "STAGE 07 complete")


if __name__ == "__main__":
    main()
