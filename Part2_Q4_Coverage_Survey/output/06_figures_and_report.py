"""
Stage 06 -- Figures, tables and the survey report
=================================================

Assembles everything the earlier stages produced into outputs a national
programme and a funding partner can use directly:

  * eight figures, styled to one system, legible in print and colour-vision safe;
  * an Excel workbook with every table on its own sheet;
  * a written survey report in markdown and Word.

Chart conventions: one measure per axis, no dual axes, recessive grid, thin
marks, direct value labels on every bar (the palette's aqua slot sits below 3:1
against the surface, so labels rather than colour carry the reading), the 95%
campaign target drawn as a reference rule wherever coverage is plotted.
"""

from __future__ import annotations

import json

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib.lines import Line2D
from matplotlib.patches import Patch

from common import (AGE_MAX_MONTHS, AGE_MIN_MONTHS, ART, AXIS, COLOR_CRITICAL, COLOR_PRIMARY,
                    COLOR_SECONDARY, COLOR_WARNING, COVERAGE_TARGET, FIGURE_DIR,
                    HOUSEHOLDS_PER_CLUSTER, INK_MUTED, INK_PRIMARY, INK_SECONDARY,
                    MOPUP_TRIGGER, PALETTE, SRC, STRATUM_NAMES, apply_chart_style, banner,
                    build_analysis_sets, get_logger, md_table)

LOG = get_logger("06_figures_and_report")
apply_chart_style()

FIGS: dict[str, str] = {}


def _save(fig, name: str, caption: str) -> None:
    path = FIGURE_DIR / f"{name}.png"
    fig.savefig(path)
    plt.close(fig)
    FIGS[name] = caption
    LOG.info("wrote figures/%s.png", name)


def _label_bars(ax, bars, values, fmt="{:.1f}", dx=0.0, dy=0.0, color=INK_PRIMARY, size=8.5):
    for b, v in zip(bars, values):
        ax.text(b.get_x() + b.get_width() / 2 + dx, b.get_height() + dy, fmt.format(v),
                ha="center", va="bottom", fontsize=size, color=color)


# --------------------------------------------------------------------------
# Figures
# --------------------------------------------------------------------------


def fig_coverage(head: pd.DataFrame) -> None:
    """Point estimate with a design-based interval, against the campaign target."""
    b = head[head.analysis_set.str.startswith("B")].copy()
    b["label"] = np.where(b.level == "National", "National", b.stratum_name)
    b = b.iloc[::-1].reset_index(drop=True)

    fig, ax = plt.subplots(figsize=(8.2, 3.6))
    colors = [COLOR_PRIMARY if r.level == "National" else PALETTE.get(r.domain, COLOR_PRIMARY)
              for r in b.itertuples()]
    y = np.arange(len(b))
    ax.hlines(y, b.ci_low_pct, b.ci_high_pct, color=colors, linewidth=2.4, alpha=0.45)
    ax.scatter(b.estimate_pct, y, s=95, color=colors, zorder=3,
               edgecolor="white", linewidth=1.4)

    ax.axvline(100 * COVERAGE_TARGET, color=COLOR_CRITICAL, linewidth=1.4, linestyle="--",
               zorder=1)
    ax.text(100 * COVERAGE_TARGET, -0.72, f"{100*COVERAGE_TARGET:.0f}% campaign target",
            color=COLOR_CRITICAL, fontsize=8.5, ha="center", va="center")

    for i, r in b.iterrows():
        ax.text(97.5, i, f"{r.estimate_pct:.1f}%  ({r.ci_low_pct:.1f}–{r.ci_high_pct:.1f})",
                va="center", ha="left", fontsize=9, color=INK_PRIMARY)

    ax.set_yticks(y, b.label)
    ax.set_ylim(-1.0, len(b) - 0.4)
    ax.set_xlim(55, 116)
    ax.set_xlabel("Campaign dose coverage, children aged 9–59 completed months (%)")
    ax.set_title("Figure 1  Weighted coverage with design-based 95% confidence intervals")
    ax.grid(axis="y", visible=False)
    ax.spines[["top", "right", "left"]].set_visible(False)
    _save(fig, "fig01_coverage_by_stratum",
          "Weighted coverage at national and stratum level. Intervals are design-based "
          "(Taylor linearisation, logit scale) and account for stratification, clustering and "
          "unequal weights. Falsified clusters excluded.")


def fig_design_effect(head: pd.DataFrame) -> None:
    """What the design cost: analysed sample against effective sample."""
    b = head[head.analysis_set.str.startswith("B")].copy()
    b["label"] = np.where(b.level == "National", "National", b.stratum_name)

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10.2, 3.9), width_ratios=[1.15, 1])
    x = np.arange(len(b))
    w = 0.38
    b1 = ax1.bar(x - w / 2, b.n_children, w, color=COLOR_PRIMARY, label="Children interviewed")
    b2 = ax1.bar(x + w / 2, b.n_effective, w, color=COLOR_SECONDARY,
                 label="Effective sample size")
    _label_bars(ax1, b1, b.n_children, "{:.0f}")
    _label_bars(ax1, b2, b.n_effective, "{:.0f}")
    ax1.set_xticks(x, b.label, fontsize=8.5)
    ax1.set_ylabel("Children")
    ax1.set_ylim(0, max(b.n_children) * 1.22)
    ax1.set_title("Figure 2a  Fieldwork bought less precision than it looks")
    ax1.legend(loc="upper right")
    ax1.grid(axis="x", visible=False)
    ax1.spines[["top", "right"]].set_visible(False)

    # Grouped rather than stacked: Kudama's total design effect is BELOW its
    # unequal-weight component, i.e. its clustering contribution is negative, and
    # a stacked bar cannot represent that honestly.
    kish = b.deff_kish_weights.to_numpy()
    total = b.deff.to_numpy()
    p1 = ax2.bar(x - w / 2, total, w, color=COLOR_PRIMARY, label="Total design effect")
    p2 = ax2.bar(x + w / 2, kish, w, color=COLOR_SECONDARY,
                 label="Unequal weights alone (Kish)")
    _label_bars(ax2, p1, total, "{:.2f}")
    _label_bars(ax2, p2, kish, "{:.2f}")
    ax2.axhline(1.0, color=AXIS, linewidth=1.1, zorder=1)   # DEFF = 1: no design penalty
    ax2.set_xticks(x, b.label, fontsize=8.5)
    ax2.set_ylabel("Design effect")
    ax2.set_ylim(0, max(total) * 1.34)
    ax2.set_title("Figure 2b  Where the design effect comes from")
    ax2.legend(loc="upper left")
    ax2.grid(axis="x", visible=False)
    ax2.spines[["top", "right"]].set_visible(False)

    fig.tight_layout()
    _save(fig, "fig02_design_effect",
          "Left: children interviewed against the effective sample size implied by the design "
          "effect. Right: the total design effect beside the part attributable to unequal "
          "weights alone; the gap between the two is the cost of clustering, and the rule at "
          "1.0 is where a design costs nothing. In Kudama the "
          "total sits below the weighting component, i.e. its clusters vary less between "
          "themselves than binomial sampling would predict, so clustering costs it nothing.")


def fig_cluster_caterpillar(set_a: pd.DataFrame, dropped: list[str]) -> None:
    """Every cluster's unweighted coverage, ordered, with the falsified ones marked."""
    cl = (set_a.groupby(["cluster_id", "stratum_code"])["vaccinated"]
          .agg(["mean", "size"]).reset_index().sort_values("mean").reset_index(drop=True))
    cl["flagged"] = cl.cluster_id.isin(dropped)

    fig, ax = plt.subplots(figsize=(10.2, 3.7))
    x = np.arange(len(cl))
    for s, color in PALETTE.items():
        m = (cl.stratum_code == s) & ~cl.flagged
        ax.scatter(x[m], 100 * cl.loc[m, "mean"], s=34, color=color,
                   label=STRATUM_NAMES[s], zorder=3, edgecolor="white", linewidth=0.8)
    m = cl.flagged
    ax.scatter(x[m], 100 * cl.loc[m, "mean"], s=64, facecolor="none",
               edgecolor=COLOR_CRITICAL, linewidth=2.0, zorder=4,
               label="Excluded: falsification screen")

    ax.axhline(100 * COVERAGE_TARGET, color=COLOR_CRITICAL, linewidth=1.3, linestyle="--")
    ax.text(1, 100 * COVERAGE_TARGET + 1.4, f"{100*COVERAGE_TARGET:.0f}% campaign target",
            color=COLOR_CRITICAL, fontsize=8.5)
    ax.annotate("nine clusters, one interviewer,\nnot one unvaccinated child",
                xy=(len(cl) - 10.5, 101.4), xytext=(len(cl) - 33, 104.5),
                fontsize=8.8, color=COLOR_CRITICAL, ha="left", va="center",
                arrowprops=dict(arrowstyle="->", color=COLOR_CRITICAL, linewidth=1.0))
    ax.set_xlabel("Clusters, ordered by measured coverage")
    ax.set_ylabel("Coverage (%)")
    ax.set_ylim(35, 111)
    ax.set_xticks([])
    ax.set_title("Figure 3  Coverage varies more between clusters than any single number shows")
    ax.legend(loc="lower right", ncols=2)
    ax.grid(axis="x", visible=False)
    ax.spines[["top", "right", "bottom"]].set_visible(False)
    _save(fig, "fig03_cluster_coverage",
          "Unweighted coverage in each of the 90 clusters. The nine ringed clusters report "
          "100% coverage with no variation whatever and were all worked by the same "
          "interviewer. Genuine cluster coverage runs from 40% to 96%, which is the spread the "
          "design effect is measuring.")


def fig_age_heaping(age: pd.DataFrame) -> None:
    fig, ax = plt.subplots(figsize=(10.2, 3.5))
    colors = np.where(age.is_year_anchor, COLOR_SECONDARY, COLOR_PRIMARY)
    ax.bar(age.age_months, age.n_children, color=colors, width=0.75)
    exp = age.expected_share_pct.iloc[0] / 100 * age.n_children.sum()
    ax.axhline(exp, color=INK_MUTED, linewidth=1.3, linestyle="--")
    ax.text(AGE_MIN_MONTHS + 1, exp + 6, "expected under a smooth age distribution",
            fontsize=8.5, color=INK_MUTED, ha="left")
    for a in age.loc[age.is_year_anchor, "age_months"]:
        n = int(age.loc[age.age_months == a, "n_children"].iloc[0])
        ax.text(a, n + 5, f"{n}", ha="center", fontsize=8.5, color=COLOR_SECONDARY)
    top = int(age.loc[age.age_months == AGE_MAX_MONTHS, "n_children"].iloc[0])
    ax.annotate(f"{top} at the eligibility ceiling:\nthe 60-month heap, kept in scope",
                xy=(AGE_MAX_MONTHS - 0.4, top + 3), xytext=(AGE_MAX_MONTHS - 15, top + 42),
                fontsize=8.5, color=INK_SECONDARY, ha="left",
                arrowprops=dict(arrowstyle="->", color=INK_MUTED, linewidth=1.0))
    ax.set_xlabel("Reported age (completed months)")
    ax.set_ylabel("Children")
    ax.set_xlim(AGE_MIN_MONTHS - 1, AGE_MAX_MONTHS + 1)
    ax.set_title("Figure 4  Age is reported in whole years, then converted")
    ax.legend(handles=[Patch(color=COLOR_SECONDARY, label="Whole-year ages (12, 24, 36, 48)"),
                       Patch(color=COLOR_PRIMARY, label="All other ages")], loc="upper left")
    ax.grid(axis="x", visible=False)
    ax.spines[["top", "right"]].set_visible(False)
    _save(fig, "fig04_age_heaping",
          "34.0% of children are reported at exactly 12, 24, 36 or 48 months against 7.8% "
          "expected. No re-analysis can recover the true ages; only a different instrument can.")


def fig_interviewers(iv: pd.DataFrame) -> None:
    fig, ax = plt.subplots(figsize=(7.6, 4.2))
    ok = iv[iv.screen_outcome != "EXCLUDE"]
    bad = iv[iv.screen_outcome == "EXCLUDE"]
    ax.scatter(ok.median_duration_min, 100 * ok.reported_coverage, s=70, color=COLOR_PRIMARY,
               edgecolor="white", linewidth=1.2, zorder=3, label="Passes the screen")
    ax.scatter(bad.median_duration_min, 100 * bad.reported_coverage, s=150,
               color=COLOR_CRITICAL, marker="X", zorder=4,
               label="Fails all four screen rules")
    for r in bad.itertuples():
        ax.annotate(f"{r.interviewer_id}\n{r.n_clusters} clusters, {r.n_children} children\n"
                    f"card seen for {100*r.card_seen_rate:.0f}% of them",
                    xy=(r.median_duration_min, 100 * r.reported_coverage),
                    xytext=(r.median_duration_min + 3.2, 100 * r.reported_coverage - 5.5),
                    fontsize=8.5, color=COLOR_CRITICAL,
                    arrowprops=dict(arrowstyle="->", color=COLOR_CRITICAL, linewidth=1.0))
    # Seventeen interviewers sit in a tight cloud, so labels are placed on
    # alternating sides and nudged vertically until they stop colliding.
    placed: list[tuple[float, float]] = []
    for i, r in enumerate(ok.sort_values("reported_coverage").itertuples()):
        side = 1 if i % 2 == 0 else -1
        tx, ty = r.median_duration_min + 0.45 * side, 100 * r.reported_coverage
        step = 0
        while any(abs(tx - px) < 1.6 and abs(ty - py) < 0.95 for px, py in placed):
            step += 1
            ty = 100 * r.reported_coverage + (0.95 * ((step + 1) // 2) * (1 if step % 2 else -1))
            if step > 8:
                break
        placed.append((tx, ty))
        ax.text(tx, ty, r.interviewer_id, ha="left" if side > 0 else "right", va="center",
                fontsize=7.4, color=INK_MUTED)
    ax.set_xlabel("Median completed-interview duration (minutes)")
    ax.set_ylabel("Reported coverage (%)")
    ax.set_xlim(0, 27)
    ax.set_ylim(72, 103)
    ax.set_title("Figure 5  One interviewer is not on the same scale as the other seventeen")
    ax.legend(loc="upper right")
    ax.spines[["top", "right"]].set_visible(False)
    _save(fig, "fig05_interviewer_screen",
          "Each point is one interviewer. Seventeen cluster tightly around a 20-minute "
          "interview and 75-89% coverage. The eighteenth completes interviews in under four "
          "minutes and reports every child vaccinated.")


def fig_source(src_strat: pd.DataFrame, meta: dict) -> None:
    piv = src_strat.pivot(index="stratum_name", columns="quantity", values="estimate_pct")
    order = ["Bansara State", "Kudama State", "Zaruwa State"]
    piv = piv.reindex(order)
    doc = piv["Documented coverage (card over all)"]
    head = piv["Headline coverage"]
    recall_part = head - doc

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10.4, 4.0), width_ratios=[1, 1])
    x = np.arange(len(piv))
    b1 = ax1.bar(x, doc, 0.55, color=COLOR_PRIMARY, label="Card-confirmed (documented)")
    b2 = ax1.bar(x, recall_part, 0.55, bottom=doc, color=COLOR_SECONDARY,
                 label="Caregiver recall only", linewidth=2, edgecolor="#fcfcfb")
    for xi, d_, r_ in zip(x, doc, recall_part):
        ax1.text(xi, d_ / 2, f"{d_:.1f}", ha="center", va="center", fontsize=9, color="white")
        ax1.text(xi, d_ + r_ / 2, f"{r_:.1f}", ha="center", va="center", fontsize=9,
                 color="white")
        ax1.text(xi, d_ + r_ + 1.2, f"{d_ + r_:.1f}%", ha="center", fontsize=9,
                 color=INK_PRIMARY)
    ax1.axhline(100 * COVERAGE_TARGET, color=COLOR_CRITICAL, linewidth=1.3, linestyle="--")
    ax1.text(len(piv) - 0.55, 100 * COVERAGE_TARGET + 1.5,
             f"{100*COVERAGE_TARGET:.0f}% target", color=COLOR_CRITICAL, fontsize=8.5,
             ha="right")
    ax1.set_xticks(x, [s.replace(" State", "") for s in piv.index])
    ax1.set_ylabel("Coverage (%)")
    ax1.set_ylim(0, 122)
    ax1.set_title("Figure 6a  How much of the estimate is documented")
    ax1.legend(loc="upper left", fontsize=8.2)
    ax1.grid(axis="x", visible=False)
    ax1.spines[["top", "right"]].set_visible(False)

    sw = pd.read_csv(ART["src_sensitivity"])
    ax2.plot(sw.recall_overreport_pct, sw.adjusted_coverage_pct, color=COLOR_PRIMARY,
             marker="o", markersize=5, zorder=3)
    ax2.fill_between(sw.recall_overreport_pct, sw.ci_low_pct, sw.ci_high_pct,
                     color=COLOR_PRIMARY, alpha=0.14, linewidth=0)
    ax2.axhline(100 * COVERAGE_TARGET, color=COLOR_CRITICAL, linewidth=1.3, linestyle="--")
    ax2.text(30, 100 * COVERAGE_TARGET + 1.0, f"{100*COVERAGE_TARGET:.0f}% target",
             color=COLOR_CRITICAL, fontsize=8.5, ha="right")
    ax2.axhline(100 * MOPUP_TRIGGER, color=COLOR_WARNING, linewidth=1.3, linestyle=":")
    ax2.text(30, 100 * MOPUP_TRIGGER + 1.0, f"{100*MOPUP_TRIGGER:.0f}% mop-up trigger",
             color="#a8760f", fontsize=8.5, ha="right")
    for r in sw.itertuples():
        if r.recall_overreport_pct in (0, 15, 30):
            ax2.text(r.recall_overreport_pct, r.adjusted_coverage_pct - 2.6,
                     f"{r.adjusted_coverage_pct:.1f}%", ha="center", fontsize=8.5,
                     color=INK_PRIMARY)
    ax2.set_xlabel("Assumed share of recall-reported doses that did not happen (%)")
    ax2.set_ylabel("National coverage (%)")
    ax2.set_ylim(60, 100)
    ax2.set_title("Figure 6b  The conclusion does not depend on recall being right")
    ax2.spines[["top", "right"]].set_visible(False)

    fig.tight_layout()
    _save(fig, "fig06_documented_source",
          "Left: the headline split into the part confirmed by a vaccination card and the part "
          "resting on caregiver recall alone. Right: the national estimate as the assumed "
          "recall over-report rate is swept from 0 to 30%. Coverage only ever moves further "
          "below target.")


def fig_sensitivity(sens: pd.DataFrame) -> None:
    d = sens.iloc[1:].copy().sort_values("shift_from_headline_pp")
    fig, ax = plt.subplots(figsize=(9.4, 5.0))
    y = np.arange(len(d))
    colors = np.where(d.shift_from_headline_pp < 0, COLOR_SECONDARY, COLOR_PRIMARY)
    ax.barh(y, d.shift_from_headline_pp, color=colors, height=0.62)
    ax.axvline(0, color=AXIS, linewidth=1.2)
    for yi, v, lab in zip(y, d.shift_from_headline_pp, d.estimate_pct):
        off = 0.09 if v >= 0 else -0.09
        ax.text(v + off, yi, f"{v:+.2f} pp  → {lab:.1f}%", va="center",
                ha="left" if v >= 0 else "right", fontsize=8.4, color=INK_PRIMARY)
    ax.set_yticks(y, [t[3:] if t[1] == "." else t[4:] for t in d.variant], fontsize=8.6)
    ax.set_xlim(-6.6, 3.0)
    ax.set_xlabel("Change in the national estimate against the headline of "
                  f"{sens.estimate_pct.iloc[0]:.1f}% (percentage points)")
    ax.set_title("Figure 7  Every assumption, replaced by its stated alternative")
    ax.grid(axis="y", visible=False)
    ax.spines[["top", "right", "left"]].set_visible(False)
    _save(fig, "fig07_sensitivity",
          "Each bar replaces one analytical assumption with the alternative documented in the "
          "assumption register. The largest single movement is the falsification exclusion. No "
          "variant brings the estimate within nine points of the 95% target.")


def fig_weights(hh: pd.DataFrame) -> None:
    comp = hh[hh.is_completed == 1].drop_duplicates("cluster_id")
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10.4, 3.8))

    for s, color in PALETTE.items():
        m = comp.stratum_code == s
        ax1.scatter(comp.loc[m, "listing_ratio"], comp.loc[m, "w_base"], s=42, color=color,
                    edgecolor="white", linewidth=0.8, label=STRATUM_NAMES[s], zorder=3)
    ax1.axvline(1.0, color=INK_MUTED, linestyle="--", linewidth=1.2)
    ax1.annotate("field listing = census size:\nthe design would be self-weighting here",
                 xy=(1.0, comp.w_base.max() * 0.62), xytext=(1.55, comp.w_base.max() * 0.72),
                 fontsize=8.2, color=INK_MUTED, ha="left",
                 arrowprops=dict(arrowstyle="->", color=INK_MUTED, linewidth=0.9))
    ax1.set_xlabel("Field listing ÷ 2023 census households")
    ax1.set_ylabel("Base design weight")
    ax1.set_title("Figure 8a  Where the unequal weights come from")
    ax1.legend(loc="upper left", fontsize=8.2)
    ax1.spines[["top", "right"]].set_visible(False)

    w = hh.loc[hh.is_completed == 1, "weight_final"]
    ax2.hist(w, bins=34, color=COLOR_PRIMARY, edgecolor="#fcfcfb", linewidth=0.6)
    ax2.axvline(w.median(), color=COLOR_SECONDARY, linewidth=1.8)
    ax2.text(w.median() * 1.06, ax2.get_ylim()[1] * 0.92, f"median {w.median():.0f}",
             color=COLOR_SECONDARY, fontsize=8.6)
    ax2.set_xlabel("Final household weight")
    ax2.set_ylabel("Households")
    ax2.set_title("Figure 8b  Distribution of the weight")
    ax2.grid(axis="x", visible=False)
    ax2.spines[["top", "right"]].set_visible(False)

    fig.tight_layout()
    _save(fig, "fig08_weights",
          "The base weight is a direct function of how far the field listing diverged from the "
          "census measure of size used for PPS selection. Had the two agreed, every point would "
          "sit on one horizontal line per stratum and the survey would have been "
          "self-weighting.")


def fig_nonresponse(ch: pd.DataFrame, hh: pd.DataFrame, dropped: list[str]) -> None:
    cl = ch.groupby("cluster_id")["vaccinated"].mean().rename("coverage")
    rr = (hh[hh.is_ineligible_dwelling == 0].groupby("cluster_id")["is_completed"].mean()
          .rename("rr"))
    st = hh.groupby("cluster_id")["stratum_code"].first()
    d = pd.concat([cl, rr, st], axis=1).dropna()
    d = d[~d.index.isin(dropped)]

    fig, ax = plt.subplots(figsize=(7.4, 4.2))
    for s, color in PALETTE.items():
        m = d.stratum_code == s
        ax.scatter(100 * d.loc[m, "rr"], 100 * d.loc[m, "coverage"], s=44, color=color,
                   edgecolor="white", linewidth=0.8, label=STRATUM_NAMES[s], zorder=3)
    z = np.polyfit(d.rr, d.coverage, 1)
    xs = np.linspace(d.rr.min(), d.rr.max(), 50)
    ax.plot(100 * xs, 100 * np.polyval(z, xs), color=INK_SECONDARY, linewidth=1.8,
            linestyle="--", zorder=2)
    ax.text(0.02, 0.96, "clusters that were hard to interview\nare clusters the campaign missed",
            transform=ax.transAxes, fontsize=8.8, color=INK_SECONDARY, va="top")
    ax.set_xlabel("Household response rate in the cluster (%)")
    ax.set_ylabel("Measured coverage in the cluster (%)")
    ax.set_title("Figure 9  Why the non-response is not plausibly ignorable")
    ax.legend(loc="lower right", fontsize=8.4)
    ax.spines[["top", "right"]].set_visible(False)
    _save(fig, "fig09_nonresponse",
          "Each point is one cluster, falsified clusters removed. Response rate and coverage "
          "are positively associated (Spearman rho = 0.39, p = 0.0003), so households the "
          "survey could not reach are likely to be households the campaign did not reach "
          "either. The weight adjustment corrects the between-cluster part of this; the "
          "within-cluster part is bounded instead.")


# --------------------------------------------------------------------------
# Workbook
# --------------------------------------------------------------------------


def write_workbook() -> None:
    sheets = {
        "T1 Coverage estimates": ART["est_headline"],
        "T2 Coverage by domain": ART["est_domains"],
        "T3 Sensitivity": ART["est_sensitivity"],
        "T4 Variance methods": ART["variance_check"],
        "T5 Coverage by source": ART["src_table"],
        "T6 Recall dependence": ART["src_sensitivity"],
        "T7 Weight diagnostics": ART["weight_diagnostics"],
        "T8 Weights by cluster": ART["weight_components"],
        "T9 Assumptions": ART["assumption_register"],
        "T10 Interviewers": ART["dq_interviewer"],
        "T11 Age heaping": ART["dq_age"],
        "T12 Missingness": ART["dq_missing"],
        "T13 Quality flags": ART["dq_flags"],
        "T14 Integrity ledger": ART["integrity_ledger"],
    }
    with pd.ExcelWriter(ART["workbook"], engine="openpyxl") as xl:
        for name, path in sheets.items():
            pd.read_csv(path).to_excel(xl, sheet_name=name[:31], index=False)
    LOG.info("wrote %s with %d sheets", ART["workbook"].name, len(sheets))


# --------------------------------------------------------------------------
# Word report
# --------------------------------------------------------------------------


def write_docx(md_text: str) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    in_table: list[str] = []

    def flush_table():
        nonlocal in_table
        if not in_table:
            return
        rows = [[c.strip() for c in r.strip().strip("|").split("|")] for r in in_table]
        rows = [r for r in rows if not all(set(c) <= set("-: ") for c in r)]
        if not rows:
            in_table = []
            return
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Light Grid Accent 1"
        for i, row in enumerate(rows):
            for j, cell in enumerate(row[:len(rows[0])]):
                para = t.cell(i, j).paragraphs[0]
                run = para.add_run(cell.replace("**", ""))
                run.font.size = Pt(8)
                if i == 0:
                    run.bold = True
        doc.add_paragraph()
        in_table = []

    for line in md_text.splitlines():
        s = line.rstrip()
        if s.startswith("|"):
            in_table.append(s)
            continue
        flush_table()
        if s.startswith("!["):
            alt = s[2:s.index("]")]
            path = s[s.index("(") + 1:s.rindex(")")]
            p = FIGURE_DIR / path.split("/")[-1]
            if p.exists():
                doc.add_picture(str(p), width=Inches(6.3))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph(alt)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].font.size = Pt(8)
                cap.runs[0].font.italic = True
                cap.runs[0].font.color.rgb = RGBColor(0x52, 0x51, 0x4E)
            continue
        if s.startswith("#"):
            level = len(s) - len(s.lstrip("#"))
            doc.add_heading(s.lstrip("# ").replace("**", ""), level=min(level, 4))
            continue
        if s.startswith(("- ", "* ")):
            doc.add_paragraph(s[2:].replace("**", ""), style="List Bullet")
            continue
        if s[:2].isdigit() and s[2:4] == ". ":
            doc.add_paragraph(s[4:].replace("**", ""), style="List Number")
            continue
        if s.startswith("> "):
            p = doc.add_paragraph(s[2:].replace("**", ""))
            p.runs[0].italic = True
            continue
        if s.strip() in ("", "---"):
            continue
        doc.add_paragraph(s.replace("**", ""))
    flush_table()
    doc.save(ART["final_docx"])
    LOG.info("wrote %s", ART["final_docx"].name)


# --------------------------------------------------------------------------


def main() -> None:
    banner(LOG, "STAGE 06  Figures, tables and the survey report")

    hh = pd.read_csv(ART["hh_weighted"])
    ch = pd.read_csv(ART["child_weighted"])
    head = pd.read_csv(ART["est_headline"])
    dom = pd.read_csv(ART["est_domains"])
    sens = pd.read_csv(ART["est_sensitivity"])
    chk = pd.read_csv(ART["variance_check"])
    src = pd.read_csv(ART["src_table"])
    age = pd.read_csv(ART["dq_age"])
    iv = pd.read_csv(ART["dq_interviewer"])
    flags = pd.read_csv(ART["dq_flags"])
    diag = pd.read_csv(ART["weight_diagnostics"])
    meta = json.loads(ART["est_headline"].with_suffix(".meta.json").read_text())
    smeta = json.loads(ART["src_sensitivity"].with_suffix(".meta.json").read_text())

    set_a, set_b, excl, dropped, _ = build_analysis_sets(ch, hh)

    fig_coverage(head)
    fig_design_effect(head)
    fig_cluster_caterpillar(set_a, dropped)
    fig_age_heaping(age)
    fig_interviewers(iv)
    fig_source(src[src.stratum_code != "National"], smeta)
    fig_sensitivity(sens)
    fig_weights(hh)
    fig_nonresponse(ch[ch.analysis_eligible], hh, dropped)

    write_workbook()

    md = build_report(head, dom, sens, chk, src, flags, diag, iv, meta, smeta, excl, dropped, hh)
    ART["final_report"].write_text(md, encoding="utf-8")
    LOG.info("wrote %s", ART["final_report"].name)
    write_docx(md)

    banner(LOG, "STAGE 06 complete")


def alloc_table(nat) -> str:
    """
    Reallocate a fixed household budget between cluster count and take per
    cluster, holding the measured intra-cluster correlation constant.

        deff = 1 + (b - 1) * rho          b = eligible children per cluster
        n_eff = n_children / deff

    The point of the table is that the last few households in a cluster are the
    cheapest to drop and the most expensive to keep.
    """
    rho = float(nat.icc_implied)
    budget = int(nat.n_clusters) * HOUSEHOLDS_PER_CLUSTER
    kids_per_hh = float(nat.n_children) / budget

    rows = []
    for take in (30, 20, 16, 12, 8, 6):
        clusters = budget // take
        b = take * kids_per_hh
        deff = 1 + (b - 1) * rho
        n_children = clusters * b
        n_eff = n_children / deff
        rows.append({
            "Households per cluster": take,
            "Clusters affordable": clusters,
            "Children interviewed": round(n_children),
            "Design effect": round(deff, 2),
            "Effective sample": round(n_eff),
            "Note": "design as executed" if take == HOUSEHOLDS_PER_CLUSTER else "",
        })
    t = pd.DataFrame(rows)
    base = t.loc[t["Households per cluster"] == HOUSEHOLDS_PER_CLUSTER, "Effective sample"].iloc[0]
    t["Change in effective sample"] = [f"{100*(v/base - 1):+.0f}%" for v in t["Effective sample"]]
    return md_table(t[["Households per cluster", "Clusters affordable", "Children interviewed",
                       "Design effect", "Effective sample", "Change in effective sample",
                       "Note"]])


def build_report(head, dom, sens, chk, src, flags, diag, iv, meta, smeta,
                 excl, dropped, hh) -> str:
    b = head[head.analysis_set.str.startswith("B")]
    a = head[head.analysis_set.str.startswith("A")]
    nat = b[b.level == "National"].iloc[0]
    natA = a[a.level == "National"].iloc[0]
    st = {r.domain: r for r in b[b.level == "Stratum"].itertuples()}
    natdiag = diag[diag.scope == "National"].iloc[0]
    nsrc = src[src.stratum_code == "National"]

    # Level and domain are merged into one column: ten columns of numbers will
    # not sit on an A4 page without the headers wrapping into nonsense.
    main_tbl = b[["level", "stratum_name", "estimate_pct", "ci_low_pct", "ci_high_pct",
                  "se_pct", "n_children", "n_clusters", "deff", "n_effective"]].copy()
    main_tbl["stratum_name"] = np.where(main_tbl.level == "National", "National (all strata)",
                                        main_tbl.stratum_name)
    main_tbl = main_tbl.drop(columns=["level"])
    main_tbl.columns = ["Domain", "Coverage %", "95% CI low", "95% CI high",
                        "SE (pp)", "Children", "Clusters", "DEFF", "Effective n"]

    def f(x, dp=1):
        return f"{x:.{dp}f}"

    L = []
    A = L.append

    A("# Post-campaign coverage survey, May 2026")
    A("## Survey report for the national programme and the funding partner")
    A("")
    A("*Three states, stratified two-stage cluster design with probability-proportional-to-size "
      "selection of enumeration areas. All estimates weighted; all confidence intervals "
      "design-based.*")
    A("")
    A("---")
    A("")
    A("## 1. Headline findings")
    A("")
    A(f"**National campaign dose coverage among children aged 9 to 59 completed months is "
      f"{f(nat.estimate_pct)}% (95% CI {f(nat.ci_low_pct)}–{f(nat.ci_high_pct)}).** That is "
      f"{f(100*COVERAGE_TARGET - nat.estimate_pct)} percentage points below the "
      f"{100*COVERAGE_TARGET:.0f}% campaign target, and the entire confidence interval lies "
      f"below it.")
    A("")
    A(f"| | Coverage | 95% CI | Children | Clusters | DEFF | Effective n |")
    A("|---|---|---|---|---|---|---|")
    A(f"| **National** | **{f(nat.estimate_pct)}%** | "
      f"{f(nat.ci_low_pct)}–{f(nat.ci_high_pct)} | {int(nat.n_children):,} | "
      f"{int(nat.n_clusters)} | {f(nat.deff, 2)} | {nat.n_effective:.0f} |")
    for code in ["ST01", "ST02", "ST03"]:
        r = st[code]
        A(f"| {STRATUM_NAMES[code]} | {f(r.estimate_pct)}% | "
          f"{f(r.ci_low_pct)}–{f(r.ci_high_pct)} | {int(r.n_children):,} | "
          f"{int(r.n_clusters)} | {f(r.deff, 2)} | {r.n_effective:.0f} |")
    A("")
    A("Four findings carry the report:")
    A("")
    A(f"1. **Coverage missed the target everywhere.** Not one stratum's confidence interval "
      f"reaches {100*COVERAGE_TARGET:.0f}%, and the shortfall is not marginal: the best "
      f"performing state is {f(100*COVERAGE_TARGET - st['ST02'].estimate_pct)} points short and "
      f"the worst is {f(100*COVERAGE_TARGET - st['ST03'].estimate_pct)} points short.")
    A(f"2. **{STRATUM_NAMES['ST03']} is a different problem from the other two.** At "
      f"{f(st['ST03'].estimate_pct)}% it sits {f(st['ST02'].estimate_pct - st['ST03'].estimate_pct)} "
      f"points below {STRATUM_NAMES['ST02']}, with the lowest household response rate and the "
      f"highest design effect in the survey. All three strata sit below the "
      f"{100*MOPUP_TRIGGER:.0f}% mop-up trigger, but {STRATUM_NAMES['ST03']} is the only one "
      f"whose interval reaches down into the sixties.")
    A(f"3. **Under half the estimate is documented.** A vaccination card was seen for only "
      f"{f(nsrc.iloc[0].estimate_pct)}% of children. Card-confirmed coverage across all "
      f"children -- the hard floor -- is {f(smeta['card_confirmed_component_pct'])}%. "
      f"{100*smeta['share_of_headline_resting_on_recall']:.0f}% of the headline rests on "
      f"caregiver recall the survey has no way to verify.")
    A(f"4. **One interviewer's work is not usable.** Nine clusters and 255 children were "
      f"submitted by an interviewer whose median interview lasted "
      f"{iv.median_duration_min.min():.0f} minutes, who reported every single child vaccinated, "
      f"and who saw a card for {100*iv.card_seen_rate.min():.0f}% of them. Including that work "
      f"would have raised the national figure to {f(natA.estimate_pct)}% -- an error of "
      f"{f(natA.estimate_pct - nat.estimate_pct)} points in the direction that would have made "
      f"the campaign look better.")
    A("")
    A(f"![Figure 1. Weighted coverage with design-based 95% confidence intervals."
      f"](../figures/fig01_coverage_by_stratum.png)")
    A("")
    A("---")
    A("")
    A("## 2. The decision: is a mop-up round required?")
    A("")
    A("**Yes, and the survey is fit to support that decision.** The reasoning is worth setting "
      "out explicitly, because 'fit for purpose' depends entirely on which purpose.")
    A("")
    A(f"The decision rule is a threshold comparison. Coverage is estimated at "
      f"{f(nat.estimate_pct)}% with an upper confidence limit of {f(nat.ci_high_pct)}%, against "
      f"a target of {100*COVERAGE_TARGET:.0f}% and a mop-up trigger of "
      f"{100*MOPUP_TRIGGER:.0f}%. The gap is roughly {f(100*COVERAGE_TARGET - nat.estimate_pct, 0)} "
      f"points, which is between three and five times the width of the confidence interval. "
      f"Every one of the {len(sens)-1} sensitivity variants -- different weighting-class "
      f"choices, trimmed weights, the opposite treatment of vacant dwellings and of "
      f"non-contacts, both bounds on item missingness, the falsified clusters retained, and an "
      f"assumed 20-point coverage deficit among non-respondents -- leaves the estimate between "
      f"{f(sens.estimate_pct.min())}% and {f(sens.estimate_pct.max())}%. None of them comes "
      f"within {f(100*COVERAGE_TARGET - sens.estimate_pct.max(), 0)} points of target.")
    A("")
    A("A conclusion that survives every analytical choice available to the analyst is a robust "
      "conclusion. **The national decision does not depend on any of the survey's weaknesses.**")
    A("")
    A("What the survey is **not** fit for:")
    A("")
    A(f"- **Targeting the mop-up below stratum level.** The design was powered for three "
      f"stratum estimates. Coverage varies from 40% to 96% between clusters, so the operational "
      f"question -- which wards to mop up -- is exactly the question the sample cannot answer. "
      f"An LGA-level estimate from this survey would rest on one to three clusters.")
    A(f"- **Measuring change against a previous round.** With an effective sample of "
      f"{nat.n_effective:.0f} children nationally, the smallest difference this survey could "
      f"detect against a comparably sized prior round -- at 5% significance and 80% power, so "
      f"2.8 x sqrt(2) x SE -- is roughly {f(2.8*np.sqrt(2)*nat.se_pct, 0)} percentage points. "
      f"Anything smaller than that is not a change, it is noise.")
    A(f"- **Attributing the shortfall to a cause.** The survey measures who was missed. It does "
      f"not measure why, and the instrument collects nothing that would support a causal "
      f"reading.")
    A("")
    A("### Stratum estimates that should not be published")
    A("")
    A(f"**{STRATUM_NAMES['ST01']} should not be published as a stratum coverage estimate.** "
      f"Five of its thirty clusters -- one sixth of the stage-one sample -- were worked by the "
      f"excluded interviewer and had to be removed. The remaining twenty-five are re-weighted "
      f"by 30/25 on the assumption that the lost clusters were exchangeable with the survivors, "
      f"and that assumption is known to be false: the clusters were not lost at random, they "
      f"were lost by interviewer assignment. Removing them moved the stratum estimate from "
      f"{f(a[a.domain=='ST01'].iloc[0].estimate_pct)}% to {f(st['ST01'].estimate_pct)}%, and "
      f"nothing in the data bounds where it would have sat had those five clusters been worked "
      f"honestly. Publish the fact of the exclusion, not a number.")
    A("")
    A(f"**{STRATUM_NAMES['ST03']} should be published only as an interval, never as a point "
      f"estimate.** Its design effect of {f(st['ST03'].deff, 2)} is the highest in the survey: "
      f"{int(st['ST03'].n_children):,} children carry the information of "
      f"{st['ST03'].n_effective:.0f}, and the confidence interval is "
      f"{f(st['ST03'].ci_width_pp)} points wide. Three of its clusters were also falsified and "
      f"removed, and it has the lowest household response rate in the survey. The defensible "
      f"public statement is '{STRATUM_NAMES['ST03']} coverage is below "
      f"{f(st['ST03'].ci_high_pct, 0)}%', which is true and decision-relevant. "
      f"'{f(st['ST03'].estimate_pct)}%' implies a precision the survey does not have.")
    A("")
    A(f"**{STRATUM_NAMES['ST02']} can be published as it stands**, with its interval. It lost "
      f"one cluster to the exclusion, its response rate is the highest in the survey, and its "
      f"design effect is below one -- its clusters are unusually homogeneous, so the "
      f"stratification and clustering cost it nothing.")
    A("")
    A("**No estimate below stratum level should be published from this survey at all** -- not "
      "by LGA, not by ward, and not by wealth quintile. The wealth quintile in particular is "
      "assigned by an asset module whose inputs were not released with the data, so it cannot "
      "be audited or recomputed, and it is observed only for households that responded.")
    A("")
    A("---")
    A("")
    A("## 3. Design, weights and precision")
    A("")
    A("### 3.1 What the weights do")
    A("")
    A("```")
    A("pi_1i    = n_h * M_i / M_h          stage one: systematic PPS on 2023 census households")
    A("pi_2k|i  = m_i / L_i                stage two: SRS of 20 from the fresh field listing")
    A("w        = f_i / (pi_1i * pi_2k|i)  with f_i the within-cluster non-response adjustment")
    A("```")
    A("")
    A(f"The single most consequential feature of this survey is that `M_i`, the census household "
      f"count used to select clusters, and `L_i`, the field listing used to select households, "
      f"are different numbers. The listing is a median 1.09 times the census count and ranges "
      f"from 0.50 to 3.77 times it. **The design is therefore not self-weighting**, the base "
      f"weight varies eight-fold across clusters, and an unweighted analysis would have been "
      f"wrong. It would have reported {f(sens[sens.variant.str.startswith('1.')].iloc[0].estimate_pct)}% "
      f"instead of {f(nat.estimate_pct)}%, and -- far more seriously -- would have reported a "
      f"standard error {f(chk.se_inflation_vs_srs.iloc[1], 2)} times too small, giving a "
      f"confidence interval {100*(1 - 1/chk.se_inflation_vs_srs.iloc[1]):.0f}% narrower than the "
      f"design justifies.")
    A("")
    A(f"![Figure 8. The listing gap is the source of the unequal weights."
      f"](../figures/fig08_weights.png)")
    A("")
    A("### 3.2 Response and coverage of the sample")
    A("")
    A(f"| Disposition of the selected sample | Households |")
    A("|---|---|")
    A(f"| Selected | {len(hh):,} |")
    A(f"| Vacant dwellings (ineligible) | {int(hh.is_ineligible_dwelling.sum())} |")
    A(f"| Refused | {int(hh.result_of_visit.eq('Refused').sum())} |")
    A(f"| Not contacted after three visits | "
      f"{int(hh.result_of_visit.eq('No eligible respondent after 3 visits').sum())} |")
    A(f"| Completed | {int(hh.is_completed.sum()):,} |")
    A(f"| **Response rate among eligible dwellings** | "
      f"**{100*hh.is_completed.sum()/(1-hh.is_ineligible_dwelling).sum():.1f}%** |")
    A("")
    A("### 3.3 Design effect and what it means for the next round")
    A("")
    A(f"The national design effect is **{f(nat.deff, 2)}**: "
      f"{int(nat.n_children):,} children interviewed under this design carry the information of "
      f"**{nat.n_effective:.0f}** children under simple random sampling. "
      f"{100*(1-nat.n_effective/nat.n_children):.0f}% of the fieldwork bought no precision at "
      f"all. The decomposition tells the programme what to do about it:")
    A("")
    A(f"- **{f(natdiag.deff_kish_children, 2)}** of it is unequal weighting, caused entirely by "
      f"the listing gap. This is fixable at design time and costs nothing to fix: select stage "
      f"one on a current listing, or re-list before selection, and the two stages cancel again.")
    A(f"- The rest is clustering, implying an intra-cluster correlation of about "
      f"**{f(nat.icc_implied, 3)}** at {nat.n_children/nat.n_clusters:.1f} children per cluster. "
      f"Vaccination status clusters because campaign teams either reached a settlement or did "
      f"not.")
    A("")
    A(f"**Recommendation for the next round: more clusters, fewer households in each.** The "
      f"table below holds the fieldwork budget fixed at the "
      f"{int(nat.n_clusters)*HOUSEHOLDS_PER_CLUSTER:,} households actually visited in the "
      f"analysed clusters, and reallocates them between cluster count and take per cluster at "
      f"the measured rho of {f(nat.icc_implied, 3)} and "
      f"{nat.n_children/(nat.n_clusters*HOUSEHOLDS_PER_CLUSTER):.2f} eligible children per "
      f"household.")
    A("")
    A(alloc_table(nat))
    A("")
    A(f"Cutting the take from 20 households to 12 raises the effective sample by about a "
      f"quarter for the same number of household visits, because at rho = "
      f"{f(nat.icc_implied, 3)} the twentieth household in a cluster carries much less new "
      f"information than the first household in a new one. Going below about 12 is where the "
      f"table stops being a fair guide: it costs the household visit only, whereas a real "
      f"budget also pays to travel to a cluster and to list it, and those fixed per-cluster "
      f"costs are what set the practical floor. Twelve is the recommendation because it is "
      f"where the precision gain is still large and the per-cluster overhead is still "
      f"affordable. {STRATUM_NAMES['ST03']}, whose rho is roughly "
      f"{f(st['ST03'].icc_implied, 2)}, gains the most from the reallocation and should also be "
      f"deliberately oversampled: it is the stratum the programme most needs a usable number "
      f"for, and it is the one this design served worst.")
    A("")
    A(f"![Figure 2. Design effect and effective sample size."
      f"](../figures/fig02_design_effect.png)")
    A("")
    A(f"![Figure 3. Coverage varies more between clusters than any single number shows."
      f"](../figures/fig03_cluster_coverage.png)")
    A("")
    A("### 3.4 Variance method")
    A("")
    A("Standard errors are Taylor-linearised under the ultimate-cluster approximation, "
      "accumulated within stratum, on 78 degrees of freedom nationally. They were validated "
      f"against a Rao-Wu-Yue rescaled bootstrap ({int(chk.bootstrap_reps.iloc[1]):,} replicates, "
      f"clusters resampled within stratum); the two agree to within "
      f"{100*abs(chk.taylor_vs_bootstrap_ratio.iloc[1]-1):.1f}%.")
    A("")
    vc = chk[["analysis_set", "estimate_pct", "se_taylor_pct", "se_bootstrap_pct",
              "se_naive_srs_pct", "se_inflation_vs_srs", "ci_taylor_pct"]].copy()
    vc.columns = ["Analysis set", "Coverage %", "SE, Taylor (pp)", "SE, bootstrap (pp)",
                  "SE if treated as SRS (pp)", "Design-based ÷ SRS", "95% CI, Taylor"]
    A(md_table(vc, {"Coverage %": ".2f", "SE, Taylor (pp)": ".3f", "SE, bootstrap (pp)": ".3f",
                    "SE if treated as SRS (pp)": ".3f", "Design-based ÷ SRS": ".2f"}))
    A("")
    A("---")
    A("")
    A("## 4. Coverage by documented source")
    A("")
    sq = nsrc[["quantity", "estimate_pct", "ci_low_pct", "ci_high_pct", "n_children"]].copy()
    sq.columns = ["Quantity", "Estimate %", "95% CI low", "95% CI high", "Children"]
    A(md_table(sq, {"Estimate %": ".1f", "95% CI low": ".1f", "95% CI high": ".1f"}))
    A("")
    A(f"The headline decomposes exactly into **{f(smeta['card_confirmed_component_pct'])} points "
      f"carried by a written record** and **{f(smeta['recall_component_pct'])} points carried by "
      f"caregiver recall alone**. Each 10 percentage points of assumed recall over-reporting "
      f"removes {f(smeta['pp_per_10pct_recall_overreport'])} points from the national estimate. "
      f"The direction of the conclusion never changes across the whole sweep -- recall error can "
      f"only make the campaign look worse.")
    A("")
    A("**A caution that matters for interpretation.** The 19-point gap between card-confirmed "
      "and recall-based coverage is *not* a measurement of recall over-reporting, and must not "
      "be presented as one. Children whose caregivers produce a card are self-selected: card "
      "retention travels with routine-immunisation contact, household stability and proximity "
      "to a facility, all of which independently predict being reached by a campaign. Part of "
      "that gap is a real coverage difference between the two groups and part is recall error. "
      "**This survey cannot separate them**, because the instrument never asks both questions "
      "of the same child.")
    A("")
    A(f"![Figure 6. Documented versus recalled coverage, and the recall sweep."
      f"](../figures/fig06_documented_source.png)")
    A("")
    A("---")
    A("")
    A("## 5. Data quality")
    A("")
    A(md_table(flags[["flag", "severity", "affects", "fixable_by"]]))
    A("")
    A("### 5.1 Falsification")
    A("")
    A(f"Four screening rules were declared before the data were examined. One interviewer fails "
      f"all four. Their median interview lasted {iv.median_duration_min.min():.0f} minutes "
      f"against a survey median of {iv.median_duration_min.median():.0f}; they reported "
      f"{100*iv.reported_coverage.max():.0f}% coverage with not one unvaccinated child across "
      f"nine clusters; they saw a card for {100*iv.card_seen_rate.min():.0f}% of children "
      f"against 46% for everyone else; and they completed "
      f"{100*iv.completion_rate.max():.0f}% of assigned households against a survey average of "
      f"{100*hh.is_completed.mean():.0f}%. The fieldwork log corroborates it independently: two "
      f"clusters and 40 households attempted in a single day.")
    A("")
    A(f"Note what the controls did **not** catch. Seven of that interviewer's eight days were "
      f"GPS verified. They were where they were supposed to be. GPS proves presence, not that "
      f"an interview happened. Only {100*pd.read_csv(SRC['fieldwork']).supervisor_spot_check.eq('Yes').mean():.0f}% "
      f"of interviewer-days carried a supervisor spot check.")
    A("")
    A(f"![Figure 5. One interviewer is not on the same scale as the other seventeen."
      f"](../figures/fig05_interviewer_screen.png)")
    A("")
    A("### 5.2 Age heaping")
    A("")
    A("34.0% of children are reported at exactly 12, 24, 36 or 48 months of age against 7.8% "
      "expected under a smooth distribution -- a 4.3-fold excess, and a Myers-type index of "
      "26.7. Caregivers report age in whole years and the interviewer converts. This does not "
      "bias the headline, which does not condition on age, but it contaminates every "
      "age-disaggregated figure and it puts real weight on the wrong side of the 9-month "
      "eligibility boundary.")
    A("")
    A("A second signature sits at the other end. **92 children are reported at exactly 59 "
      "months**, three times the average of the five months below it. Sixty months is outside "
      "the definition, so the whole-year heap that would have landed there lands instead on the "
      "last age that keeps a child eligible. That is a problem with the *eligible population*, "
      "not only with the age variable: an unknown share of those children are five years old, "
      "were never in scope, and are old enough to have been reached by an earlier round.")
    A("")
    A(f"![Figure 4. Age is reported in whole years, then converted."
      f"](../figures/fig04_age_heaping.png)")
    A("")
    A("### 5.3 Missingness, and whether it is ignorable")
    A("")
    A("**Item missingness (44 children, 1.9%): plausibly ignorable.** It is spread across 15 "
      "interviewers and all three strata and is not concentrated in any subgroup. Counting all "
      "44 as vaccinated gives 81.7%; counting all as unvaccinated gives 79.6%. The headline "
      "sits inside a two-point band whatever the truth is, so the assumption carries no weight. "
      "Handled by complete-case analysis inside the weighted estimator, with both bounds "
      "reported.")
    A("")
    A(f"**Unit non-response ({100*(1-hh.is_completed.sum()/(1-hh.is_ineligible_dwelling).sum()):.0f}% "
      f"of eligible households): not plausibly ignorable.** At cluster level the response rate "
      f"and measured coverage move together (Spearman rho = 0.39, p = 0.0003). The clusters "
      f"that were hard to interview are the clusters the campaign missed -- absent households, "
      f"mobile populations and inaccessible settlements defeat the survey team and the "
      f"vaccination team alike. A mechanism operating between clusters almost certainly operates "
      f"within them. Whether it does cannot be tested, because no outcome was recorded for a "
      f"single non-responding household.")
    A("")
    A("**What was done.** The weight adjustment is made within cluster, the finest weighting "
      "class the data support, which absorbs the between-cluster part. The within-cluster part "
      "cannot be absorbed, so it is bounded: an assumed 10-point coverage deficit among "
      "non-respondents puts the national estimate at 79.2%, a 20-point deficit at 77.4%. Those "
      "are the figures a decision maker should treat as the plausible lower end.")
    A("")
    A(f"![Figure 9. Why the non-response is not plausibly ignorable."
      f"](../figures/fig09_nonresponse.png)")
    A("")
    A("---")
    A("")
    A("## 6. Sensitivity analysis")
    A("")
    sv = sens[["variant", "estimate_pct", "ci_low_pct", "ci_high_pct",
               "shift_from_headline_pp"]].copy()
    sv.columns = ["Variant", "Coverage %", "95% CI low", "95% CI high", "Shift (pp)"]
    A(md_table(sv, {"Coverage %": ".1f", "95% CI low": ".1f", "95% CI high": ".1f",
                    "Shift (pp)": "+.2f"}))
    A("")
    A(f"![Figure 7. Every assumption, replaced by its stated alternative."
      f"](../figures/fig07_sensitivity.png)")
    A("")
    A("---")
    A("")
    A("## 7. Limitations: what better analysis could fix, and what only a better instrument could")
    A("")
    A("This distinction determines where the next round's money should go. Analysis is cheap "
      "and can be redone; instruments and field protocols have to be changed before fieldwork "
      "starts and cannot be retrofitted.")
    A("")
    A("### 7.1 Fixable by better analysis -- and fixed here")
    A("")
    A("| Limitation | How it was fixed | Effect on the headline |")
    A("|---|---|---|")
    A("| The design is not self-weighting, because the field listing diverges from the census "
      "measure of size | Second-stage probability computed from the field listing, not the "
      "census count | +1.6 points against the self-weighting assumption |")
    A("| Treating the sample as simple random | Ratio estimator with Taylor-linearised, "
      "stratified, ultimate-cluster variance | point estimate +0.7 points; standard error "
      f"{f(chk.se_inflation_vs_srs.iloc[1], 1)}x larger, which is the material correction |")
    A("| An entire cluster submitted twice | De-duplicated on the full record | removed 20 "
      "households and 20 children that would have been double-counted in both the estimate and "
      "the variance |")
    A("| Between-cluster differences in non-response | Within-cluster weight adjustment | "
      "-0.6 points |")
    A("| Vacant dwellings confused with refusals | Vacancy classified as ineligibility, not "
      "non-response | -0.1 points |")
    A("| Item missingness on the outcome | Bounded in both directions | +0.4 / -1.7 points |")
    A("| Children enumerated outside the eligible age range | Excluded, and the exclusion "
      "tested | +0.1 points |")
    A("| Falsified records | Pre-declared four-rule screen; nine clusters removed | -1.8 "
      "points, and the largest single correction in the analysis |")
    A("")
    A("Note the asymmetry in the last row. **Analysis could detect the falsification; it could "
      "not repair it.** Those nine clusters are gone. Their strata are represented by twenty-five "
      "and twenty-seven clusters instead of thirty, on an exchangeability assumption that is "
      "known to be false. Detection is an analytical capability; prevention is a field-protocol "
      "capability, and only the second one would have left the survey whole.")
    A("")
    A("### 7.2 Fixable only by a better instrument or a better field protocol")
    A("")
    A("| Limitation | Why no analysis can fix it | What would fix it |")
    A("|---|---|---|")
    A("| **Recall accuracy is unidentifiable.** Q3.5 is asked only when a card is seen and Q3.6 "
      "only when it is not, so no child ever has both | There is no subsample on which recall "
      "can be validated against a card. The 19-point gap between the two groups confounds recall "
      "error with genuine coverage differences between card-holding and non-card-holding "
      "households, and nothing in the data separates them | Ask both of everyone: record the "
      "caregiver's report first, then ask for the card. Thirty seconds per child, and recall "
      "accuracy becomes directly estimable on the card-holding subsample |")
    A("| **Age heaping.** 34% of ages sit on whole years | Smoothing the distribution would "
      "manufacture ages that were never observed. The heaping is in the measurement, and no "
      "estimator recovers information that was not collected | Collect date of birth from a "
      "document where one exists; use a local-events calendar to probe where it does not; make "
      "the instrument reject an age that contradicts a recorded birth date |")
    A("| **Within-cluster non-response bias.** Nothing at all is known about non-responding "
      "households | A weight adjustment can only redistribute within a class using the "
      "responders' own outcomes. It cannot know that the non-responders differ, and here there "
      "is direct evidence that they do | A minimal non-response form: for every refusal or "
      "non-contact, record dwelling type, presence of children observed, and a neighbour-proxy "
      "vaccination report. Two minutes per household turns an untestable assumption into an "
      "estimable one |")
    A("| **Card retention at 46%** caps how much of any estimate can ever be documented | "
      "Analysis cannot create records that do not exist | A programme fix (issue and replace "
      "cards, use home-based records) plus an instrument fix (record whether a card was ever "
      "issued, separately from whether it was produced today) |")
    A("| **Falsification prevention.** GPS verification confirmed presence but not "
      "interviewing | Post-hoc screening is a detection tool with a real false-negative rate. A "
      "more careful falsifier -- one who reported 88% coverage and 18-minute interviews -- would "
      "have passed every rule here | Timestamped question-level audit trails, random audio "
      "audits, and independent re-interview of 5% of households within 72 hours. Supervisory "
      "spot checks covered only 19% of interviewer-days |")
    A("| **Interview duration recorded as a whole number by the interviewer** | A "
      "self-reported duration is exactly the field a falsifier controls | Derive duration from "
      "the device's own timestamps, where it cannot be edited |")
    A("| **Precision at stratum level.** DEFF up to 3.3, effective n as low as 188 | No "
      "estimator recovers precision that the sample allocation did not buy | Reallocate: more "
      "clusters, fewer households each, and oversample the stratum expected to be worst |")
    A("| **Three clusters replaced because the original was inaccessible** | The replacement "
      "carries the probability of the cluster it replaced, which is defensible only if the two "
      "are exchangeable -- and inaccessible areas are systematically harder to vaccinate | A "
      "protocol that documents the reason, the replacement rule and the characteristics of the "
      "cluster that was lost, so that the direction of the bias can at least be signed |")
    A("| **Wealth quintile cannot be audited** -- the asset module inputs were not released | "
      "The variable cannot be recomputed or checked against its own components | Release the "
      "asset module with the data, or publish the index construction |")
    A("")
    A("### 7.3 The one-sentence version")
    A("")
    A("> Better analysis fixed the weighting, the duplication, the classification of "
      "non-response and the detection of falsified work, and it moved the headline by about two "
      "points; but the two limitations that actually bound what this survey can claim -- that "
      "recall accuracy is unidentifiable by construction, and that nothing whatever is known "
      "about the households that did not respond -- are properties of the instrument, and no "
      "amount of re-analysis will touch them.")
    A("")
    A("---")
    A("")
    A("## 8. Recommendations")
    A("")
    A("**For the current campaign**")
    A("")
    A(f"1. **Conduct a mop-up round.** Coverage is {f(nat.estimate_pct)}% against a "
      f"{100*COVERAGE_TARGET:.0f}% target, and the conclusion is robust to every analytical "
      f"choice tested.")
    A(f"2. **Prioritise {STRATUM_NAMES['ST03']}**, which is roughly "
      f"{f(st['ST02'].estimate_pct - st['ST03'].estimate_pct, 0)} points behind the best "
      f"performing state and is the only stratum clearly below the mop-up trigger.")
    A("3. **Do not use this survey to choose wards.** Use programme microplanning data, "
      "administrative coverage and the cluster-level results as corroboration, not as a "
      "sampling frame for targeting.")
    A("4. **Re-survey the nine excluded clusters** if any stratum-level number is to be "
      "published for Bansara State. They are five of that stratum's thirty clusters and no "
      "adjustment substitutes for the data.")
    A("")
    A("**For the next survey round**")
    A("")
    A("5. **Select stage one on a current listing**, or re-list before selection. This removes "
      f"the {100*(natdiag.deff_kish_children-1):.0f}% of the design effect that unequal "
      f"weighting is costing, at no fieldwork cost.")
    A("6. **Take 12 households per cluster, not 20, and raise the cluster count.** At rho = "
      f"{f(nat.icc_implied, 3)} this buys substantially more precision for the same budget.")
    A("7. **Ask the recall question of every caregiver, before asking for the card.** This is "
      "the single highest-value change to the instrument and it costs thirty seconds per child.")
    A("8. **Add a minimal non-response form.** Without it, the largest single uncertainty in "
      "this report stays untestable in the next one.")
    A("9. **Replace self-reported interview duration with device timestamps, and add random "
      "audio audit plus 5% independent re-interview.** The falsification here was caught by "
      "analysis; the next one may not be.")
    A("")
    A("---")
    A("")
    A("## Annex A. Full estimate table")
    A("")
    A(md_table(main_tbl, {"Coverage %": ".1f", "95% CI low": ".1f", "95% CI high": ".1f",
                          "SE (pp)": ".2f", "DEFF": ".2f", "Effective n": ".0f"}))
    A("")
    A("## Annex B. Coverage by domain")
    A("")
    A("*Reported for completeness and internal programme use. Sub-stratum domains are not "
      "publishable estimates -- see section 2.*")
    A("")
    dm = dom[["domain_variable", "domain_label", "estimate_pct", "ci_low_pct",
              "ci_high_pct", "n_children", "deff", "n_effective"]].copy()
    dm["domain_variable"] = dm["domain_variable"].map(
        {"stratum_code": "State", "settlement_type": "Settlement",
         "wealth_label": "Wealth quintile", "sex": "Sex", "age_group": "Age group",
         "status_source": "Source of evidence"}).fillna(dm["domain_variable"])
    dm.columns = ["Disaggregation", "Category", "Coverage %", "95% CI low", "95% CI high",
                  "Children", "DEFF", "Effective n"]
    A(md_table(dm, {"Coverage %": ".1f", "95% CI low": ".1f", "95% CI high": ".1f",
                    "DEFF": ".2f", "Effective n": ".0f"}))
    A("")
    A("## Annex C. Figures")
    A("")
    for name, cap in FIGS.items():
        A(f"**{name}** -- {cap}")
        A("")
    A("## Annex D. Files")
    A("")
    A("| File | Contents |")
    A("|---|---|")
    A("| `reports/06_survey_report_tables.xlsx` | every table in this report, one per sheet |")
    A("| `reports/01_preparation_and_validation.md` | structural validation and integrity ledger |")
    A("| `reports/02_design_weights.md` | weight derivation and the full assumption register |")
    A("| `reports/03_weighted_estimates.md` | estimates, design effects, variance validation |")
    A("| `reports/04_data_quality_assessment.md` | heaping, interviewer screen, missingness |")
    A("| `reports/05_documented_source_analysis.md` | card versus recall |")
    A("| `data/children_weighted.csv` | analysis file with every weight variant |")
    A("| `figures/` | the nine figures, at 200 dpi |")
    A("")
    return "\n".join(L)


if __name__ == "__main__":
    main()
