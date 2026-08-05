"""
Render "A critical reading of the questionnaire" as a branded Word document.

This is a synthesis for a reader who wants one artefact answering: was the paper
instrument read critically, and were its flaws handled with judgement? The
underlying detail stays in docs/04_defect_register.md (24 findings),
docs/03_coding_and_sentinels.md (15 collisions) and the generated constraint
register.
"""

from __future__ import annotations

import os
import re
import sys

import docx
from docx.shared import Inches, Pt as DPt

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import doctheme as T  # noqa: E402

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.normpath(os.path.join(HERE, "..", "docs"))
os.makedirs(OUT, exist_ok=True)

BODY = 9.4
TABLE = 8.3
LINE = 1.13

_TOK = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*`]+?\*)")


def rich(p, text, size=BODY):
    for tok in _TOK.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.font.name, r.font.bold = T.FONT, True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = T.MONO; r.font.color.rgb = T.D_BLUE
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.font.name, r.font.italic = T.FONT, True
        else:
            r = p.add_run(tok); r.font.name = T.FONT
        r.font.size = DPt(size)
    return p


def tbl(doc, header, rows, widths, size=TABLE):
    """add_table with markup stripped and sane page-break behaviour.

    Table cells are plain runs, so **bold** inside one would print its own
    asterisks. And a table that crosses a page break must repeat its header row
    and must not split a row down the middle - both of which python-docx leaves
    off by default.
    """
    t = T.add_table(doc,
                    [T.strip_md(c) for c in header] if header else None,
                    [[T.strip_md(str(c)) for c in row] for row in rows],
                    widths=widths, size=size)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for i, row in enumerate(t.rows):
        pr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")          # keep a row whole
        pr.append(cant)
        if i == 0 and header:
            hdr = OxmlElement("w:tblHeader")       # repeat the header row
            pr.append(hdr)
    return t


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = DPt(11 if level == 1 else 8)
    p.paragraph_format.space_after = DPt(3)
    return p


def para(doc, text, size=BODY, space_after=5, indent=0.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = DPt(space_after)
    p.paragraph_format.line_spacing = LINE
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    return rich(p, text, size=size)


def bullet(doc, text, size=BODY):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = DPt(3)
    p.paragraph_format.line_spacing = LINE
    return rich(p, text, size=size)


def squeeze(doc, pt=3):
    p = doc.paragraphs[-1]
    if p.text.strip():
        return
    p.paragraph_format.space_after = DPt(0)
    p.paragraph_format.space_before = DPt(0)
    (p.runs[0] if p.runs else p.add_run("")).font.size = DPt(pt)


def finding(doc, ident, title, blocks):
    """One detailed finding: id + title, then labelled paragraphs."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = DPt(9)
    p.paragraph_format.space_after = DPt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{ident}   ")
    r.font.name, r.font.size, r.font.bold = T.FONT, DPt(10.5), True
    r.font.color.rgb = T.D_BLUE
    r2 = p.add_run(title)
    r2.font.name, r2.font.size, r2.font.bold = T.FONT, DPt(10.5), True
    for lead, text in blocks:
        q = doc.add_paragraph()
        q.paragraph_format.space_after = DPt(2)
        q.paragraph_format.line_spacing = LINE
        q.paragraph_format.left_indent = Inches(0.28)
        q.paragraph_format.first_line_indent = Inches(-0.28)
        rl = q.add_run(f"{lead}  ")
        rl.font.name, rl.font.size, rl.font.bold = T.FONT, DPt(BODY), True
        rl.font.color.rgb = T.D_BLUE
        rich(q, text, size=BODY)


def build() -> str:
    doc = docx.Document()
    T.docx_setup(doc, margin=0.72)

    T.cover(
        doc,
        "A critical reading of the questionnaire",
        "Form HH/2026 · What is wrong with the paper instrument, and how each flaw was treated",
        meta_lines=[
            "Integrated Child Health and Antimicrobial Resistance Survey 2026  ·  Ethics approval BSHREC/2026/041",
            "Companion to: 04_defect_register.md (full detail) · 03_coding_and_sentinels.md · 02_constraint_register.csv",
        ],
        kicker="24 defects · 15 coding collisions · 13 resolved · 6 resolved and escalated · 5 escalated · 2 blocking",
    )

    # ------------------------------------------------------------------ 1
    h(doc, "1  The claim")
    T.callout(
        doc, "In one sentence",
        "A digital form that transcribes this questionnaire faithfully would be a working form "
        "that still permits every error the paper form permits - and one of those errors "
        "destroys the measurement it records, at the point of capture, irreversibly.",
    )
    squeeze(doc)
    para(doc,
         "The questionnaire is a competent paper instrument. It is also internally contradictory in "
         "places, has a skip instruction missing outright, asks one question that cannot be answered "
         "as printed, prints two code boxes too small for the codes they are meant to hold, and "
         "carries a sentinel value inside a measurement field where the sentinel is a perfectly "
         "ordinary measurement. None of that is unusual. All of it survives transcription.")
    para(doc,
         "This document sets out what was found and — more importantly — the rule used to decide "
         "what to fix silently, what to fix and escalate, what to escalate untouched, and what to "
         "refuse to guess at.")

    # ------------------------------------------------------------------ 2
    h(doc, "2  How each finding was dispositioned")
    para(doc,
         "The instrument carries ethics approval. Approval attaches to **the questions asked and the "
         "data collected**, not to the layout of the skip column. That single distinction decides "
         "almost every case:")
    tbl(
        doc,
        ["Disposition", "Applied when", "n"],
        [
            ["Resolved in the form",
             "The paper form is internally inconsistent and only one reading is coherent, or the fix "
             "changes how a value is captured without changing what is captured", "13"],
            ["Resolved and escalated",
             "The fix was necessary and unambiguous, but it changes field behaviour in an approved "
             "instrument, so the committee is told and it is one row to revert", "6"],
            ["Added and escalated",
             "One question was added that the instrument does not contain. Flagged for ratification", "1"],
            ["Escalated only",
             "Fixing it would change what is asked or recorded. Not mine to decide", "5"],
            ["Removed",
             "Meaningless in a digital instrument. Recorded so the omission is a decision, not an oversight", "1"],
            ["Blocking",
             "A required input does not exist. Guessing would look like an answer", "2"],
        ],
        widths=[22, 68, 10])
    squeeze(doc)
    para(doc,
         "The bar for *silent* correction is deliberately high. Two changes alter what an enumerator "
         "is **permitted to do** rather than how a value is **stored** — the consent block and the "
         "added adult-respondent check — and both are flagged for the ethics committee even though I "
         "believe both are right. Silently fixing a contradiction in an ethics-approved instrument is "
         "not automatically the correct answer.")

    # ------------------------------------------------------------------ 3
    h(doc, "3  The findings that change the data")
    para(doc, "Seven of the twenty-four. The rest are in the defect register.", space_after=3)

    finding(doc, "D-8", "4.06 stores \"not measured\" as a value that is also a real height  (Critical)",
            [("Paper",
              "Length or height is a three-digit field with one decimal — `___._ cm` — carrying "
              "`Not measured . . . 99` in the same boxes as the measurement."),
             ("Why it matters",
              "**99.0 cm is an ordinary height for a child of three to four**, well inside the 9–59 month "
              "range this module covers. Once written, a 99.0 in that column is irreducibly ambiguous: no "
              "plausibility rule, no cross-check, no reconciliation against the weight can separate a "
              "measured child from a refusal. The information is destroyed at the point of capture — this "
              "is the one defect that cannot be repaired downstream by anybody, ever."),
             ("Treated",
              "**Resolved.** A gate question asks whether the child was measured. The centimetre field is "
              "never entered when they were not, and a coded reason is stored instead. 99.0 remains "
              "available as a real measurement."),
             ("Why that",
              "It changes how the fact is stored, not what is collected, so it needs no approval. Test "
              "T-06e asserts 99.0 cm is **accepted** — the one boundary case where accepting is the "
              "correct behaviour.")])

    finding(doc, "D-10", "5.05 cannot record the failure it exists to detect  (High)",
            [("Paper",
              "Cold box temperature is printed as one digit and one decimal — `_._ °C`."),
             ("Why it matters",
              "**No reading of 10 °C or above can be written down.** A cold box left in the sun at 31 °C — "
              "precisely the event the question exists to catch — has no representation on the form. The "
              "enumerator's only options are to write something false or leave it blank."),
             ("Treated",
              "**Resolved and escalated.** Range widened to −5.0 to 40.0 °C, with a confirmation prompt "
              "outside the 2–8 °C target band. Escalated because the printed field must be widened before "
              "the paper form is used anywhere again.")])

    finding(doc, "D-1", "3.02 cannot be answered as printed  (Critical)",
            [("Paper",
              "*\"From column (7), how many children in this household are aged 9 to 59 completed "
              "months?\"* Column (7) is headed **\"Eligible for Section 4 (office use)\"** and the "
              "interviewer instruction says it *\"is completed by the office and must be left blank in "
              "the field.\"*"),
             ("Why it matters",
              "The enumerator must read a count off a column they are forbidden to fill in. 3.02 is the "
              "gate deciding whether Section 4 happens at all, so the whole child module hangs off an "
              "unanswerable question."),
             ("Treated",
              "**Resolved and escalated.** Eligibility is computed from roster ages. 3.02 is still asked "
              "independently, then reconciled against the computed count by a gate that cannot be "
              "dismissed until the two agree.")])

    finding(doc, "D-3", "5.02 has no skip instruction at all  (High)",
            [("Paper",
              "*\"Was a stool specimen obtained from this child?\"* has an **empty SKIP column**."),
             ("Why it matters",
              "5.03–5.05 (label, cold box time, temperature) apply only if a specimen was obtained; "
              "5.06–5.07 (reason none was obtained) only if it was not. As printed, both branches are "
              "open on both answers, and nothing tells the enumerator which to complete."),
             ("Treated",
              "**Resolved.** Yes opens 5.03–5.05, No opens 5.06–5.07, neither can be left incomplete. The "
              "intent is unambiguous, so this needed no escalation.")])

    finding(doc, "D-7", "2.01 = No has no consequence  (Critical)",
            [("Paper",
              "*\"Consent statement read aloud to the respondent in full?\"* — No is a permitted answer "
              "with no skip and no instruction."),
             ("Why it matters",
              "The form proceeds to 2.02 and records consent to a statement that was never read. That is "
              "not a data quality defect. It is an instrument that permits an invalid consent to be "
              "documented as valid."),
             ("Treated",
              "**Resolved and escalated.** 2.01 must be Yes to continue; the message directs the "
              "enumerator to read the statement now, or to record the visit as refused at 1.14 if the "
              "respondent will not allow it. No data are lost — a genuine refusal already has a code."),
             ("Why escalated",
              "A hard block changes field behaviour in an approved instrument. I believe it is the only "
              "defensible reading, and the committee should still ratify it.")])

    finding(doc, "D-16", "The LGA and ward code boxes are the wrong size  (High)",
            [("Paper",
              "1.02 and 1.03 each print `Code` followed by three digit boxes."),
             ("Why it matters",
              "The official codes in the supplied lookups are `LGA02` (five characters) and `W018` "
              "(four). Neither fits. As printed, the approved form **cannot record the codes it asks "
              "for**, and 120 enumerators will improvise differently. Settlement is correct at six boxes "
              "for `S01324`, which shows the error is an oversight rather than a different convention."),
             ("Treated",
              "**Resolved.** Codes are never keyed. LGA, ward and settlement are cascading selections "
              "from the supplied lookup files and the code is stored automatically.")])

    finding(doc, "D-24", "1.13 invites re-linking households that declined follow-up  (High)",
            [("Paper",
              "1.13 asks for the prior household identifier for any household visited in the last round."),
             ("Why it matters",
              "**342 of the 3,982 households** in the previous round register carry "
              "`consent_to_follow_up = no`. Nothing on the form distinguishes them. Recording the "
              "identifier links this round's child health and specimen data to households that asked "
              "not to be re-approached. This one is not visible on the questionnaire at all — it is only "
              "visible if you read the questionnaire against the data files it depends on."),
             ("Treated",
              "**Resolved and escalated.** Those rows are filtered out before the lookup reaches a "
              "device, so they cannot be selected. Escalated to the data protection lead, because the "
              "records remain re-linkable by anyone holding both datasets and that is a policy decision.")])

    # ------------------------------------------------------------------ 4
    h(doc, "4  Where a coding category collides with a non-response code")
    para(doc,
         "The notes on completion impose a global sentinel scheme — `8`/`98` for don't know, `9`/`99` "
         "for no answer, `96` for Other — on question lists that were designed independently of it. "
         "**Fifteen** places were checked; one of them is confirmed safe and listed anyway, because a "
         "register that only lists problems does not show the others were examined.")
    tbl(
        doc,
        ["Where", "The collision", "Severity"],
        [
            ["4.06 Height", "\"Not measured = 99\" in a field where 99.0 cm is a real height", "Critical"],
            ["6.01 Water", "8 = Unprotected spring vs DK; 9 = Rainwater vs no answer", "High"],
            ["6.02 Toilet", "8 = Bucket vs DK; 9 = No facility or bush vs no answer — and 9 is the open-defecation category", "High"],
            ["1.02 / 1.03 codes", "Code box width incompatible with the official code format (D-16)", "High"],
            ["4.05 Weight", "\"Not measured = 99\" inside a kilogram field whose own maximum is 99.9", "Medium"],
            ["Roster col (5)", "Age in years: 98 and 99 are attainable human ages", "Medium"],
            ["5.03 Check digit", "'X' is a letter in a field the notes treat as digit boxes", "Medium"],
            ["6.07 Assets", "Letter codes A–H; the numeric sentinel scheme does not apply. 'None of these' behaves like a sentinel", "Medium"],
            ["5.05 Temperature", "No sentinel at all; blank is indistinguishable from a genuine 0.0 °C", "Medium"],
            ["Global note", "\"Non-response codes are used throughout\" — seven questions print no such category", "Medium"],
            ["4.13 Antibiotic", "Cannot confirm the ministry list avoids 96/98/99 — the list does not exist", "Unknown"],
            ["3.01 / 3.02 counts", "Two-digit fields; 98/99 collide with counts. Latent, not actual", "Low"],
            ["1.06 / 1.07", "Three-digit fields; the notes define no sentinel for that width", "Low"],
            ["Roster col (6)", "Age in months, max 59 — sentinels sit safely above the range. **Confirmed safe**", "None"],
        ],
        widths=[18, 68, 14])
    squeeze(doc)
    para(doc,
         "**The storage rule adopted:** a sentinel is never stored in a field that carries a quantity. "
         "Every measurement is preceded by a gate; the numeric field is empty when nothing was "
         "measured and a separate coded field records why. Sentinels remain inside categorical fields "
         "**only** where the substantive list provably cannot collide — which is true of the seven "
         "yes/no questions and false of 6.01 and 6.02, where no don't-know option is offered at all. "
         "The result: `99` appears nowhere in the dataset as a value, and no numeric column mixes a "
         "measurement with a code.")

    # ------------------------------------------------------------------ 5
    h(doc, "5  What I did not fix")
    para(doc,
         "Restraint is the harder half of this. Five findings were escalated **untouched**, because "
         "resolving them either way changes what is asked or recorded:")
    for t in [
        "**D-2b** The notes claim non-response codes are used \"throughout\"; seven questions print no "
        "such category. Either the notes overstate the rule or those questions are missing categories. "
        "Deciding changes what is collected, so sentinels are offered only where printed.",
        "**D-17** The header says fieldwork is 1–30 June; the operating conditions say 14 days. The form "
        "enforces the narrower window and the ministry is asked which is authoritative.",
        "**D-14** \"Eligible\" means 9–59 months in Section 4 and 12–59 months in Section 5. Escalated as "
        "wording; in the form they are two separate computed variables that cannot be confused.",
        "**D-21** 1.12 asks about \"the October 2025 round\", but the register's visits run to 9 December "
        "2025. Escalated; the form does not constrain on the month.",
        "**D-6** 2.02's \"→ END\" carries no signing instruction where 1.14's does. Resolved "
        "conservatively — Section 7 always opens — and escalated as an inconsistency between two "
        "identical instructions in approved text.",
    ]:
        bullet(doc, t)
    para(doc,
         "**One question was added** to an approved instrument: 2.01a, *is the respondent aged 18 or "
         "over*. The form records the respondent's relationship to the head but never their age or "
         "capacity, so a 14-year-old is a valid respondent as printed. It collects one boolean, no new "
         "personal data, and is flagged for ratification. It is one row to remove.")

    # ------------------------------------------------------------------ 6
    h(doc, "6  What I refused to invent")
    para(doc,
         "Two required inputs do not exist. In both cases a plausible guess was available and would "
         "have produced a form that looked finished. Both are recorded as **blocking** instead.")
    tbl(
        doc,
        ["", "D-15  The medicine list", "D-15b  The check digit scheme"],
        [
            ["What is missing",
             "4.13 says \"record from the medicine list\". No such list exists in the questionnaire, in "
             "reference_media/, or anywhere in the pack",
             "The allocation file states the scheme but ships no example label carrying its check digit"],
            ["Why guessing is not safe",
             "The survey's core AMR variable would be coded against a scheme the ministry never "
             "approved. Drugs absent from an invented list are unrecoverable — you cannot recover a "
             "distinction the instrument never offered",
             "Two readings of \"remainder\" are current and they disagree on **90.9% of the 21,600 "
             "issued serials**. If the wrong one is implemented, the form rejects a correctly printed "
             "label 91% of the time, offline, on day one"],
            ["What was done instead",
             "A clearly-marked placeholder, and a build that fails if the eventual approved list "
             "assigns a real medicine to code 96, 98 or 99",
             "The literal reading implemented, the alternative implemented beside it, and a one-line "
             "switch. Resolvable by reading one physical label"],
        ],
        widths=[16, 42, 42])
    squeeze(doc)
    para(doc,
         "The second one is worth dwelling on, because the failure is not merely that work stops. An "
         "enumerator blocked by a form insisting a real label is wrong will change digits until one "
         "passes — producing a serial that satisfies the check digit but **does not match the tube in "
         "their hand**. That is exactly the \"specimen cannot be matched to a child record\" case the "
         "laboratory answers by discarding the specimen and sending the team back. A wrong guess here "
         "would not just halt fieldwork; it would silently corrupt the linkage the specimen section "
         "exists to protect.")

    # ------------------------------------------------------------------ 7
    h(doc, "7  How the treatment is held in place")
    para(doc,
         "Findings decay unless something enforces them. Each of these is machine-checked on every "
         "build, and each was regression-tested by reintroducing the defect it guards against:")
    tbl(
        doc,
        ["Mechanism", "What it guarantees"],
        [
            ["Constraint register generated from the form",
             "60 rules, each carrying what it prevents and where its threshold came from. A constraint "
             "added without a justification **fails the build**. The register cannot describe a rule the "
             "form does not contain"],
            ["16 of 60 thresholds labelled \"judgement\"",
             "Where the questionnaire states no bound and no source could be named, the register says so "
             "rather than implying a standard. The weight and height bounds are deliberately wide, to "
             "catch keying errors rather than clinical outliers"],
            ["97 structural checks, run over both language variants",
             "Asserts the compiled XForm contains the rule that was intended — including the ones that "
             "silently compiled wrong the first time"],
            ["87 boundary cases against the deployed expressions",
             "Read the constraint out of the compiled form and apply the candidate value, so the tests "
             "follow the form rather than restating it. T-06e asserts 99.0 cm is accepted"],
            ["459,604 check-digit cases",
             "Every transposition of two unequal digits across every issued serial is rejected — "
             "292,960 cases, none accepted"],
            ["Variant equivalence diff",
             "The two language forms are compared bind by bind; a rule fixed in one and not the other "
             "fails the build"],
        ],
        widths=[30, 70])

    para(doc,
         "**The measure of the exercise is not the count of defects found.** It is that each one has a "
         "recorded disposition, a stated reason for that disposition, and — where it was resolved — a "
         "test that fails if the resolution is ever undone.", space_after=2)

    path = os.path.join(OUT, "00_critical_reading_of_the_questionnaire.docx")
    doc.save(path)
    return path


if __name__ == "__main__":
    p = build()
    print("wrote", p, f"({os.path.getsize(p) / 1024:.0f} KB)")
