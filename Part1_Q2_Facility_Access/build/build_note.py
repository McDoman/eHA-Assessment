"""
Render reports/06_methodological_note.md as a branded Word document.

The markdown stays the source of record; this is a presentation layer for it.
Nothing under Outputs/reports is modified.
"""

from __future__ import annotations

import os
import re
import sys

import docx
from docx.shared import Pt as DPt

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import doctheme as T  # noqa: E402

HERE = os.path.dirname(os.path.abspath(__file__))
# NB: the repo already has an "Outputs" directory and Windows paths are
# case-insensitive, so a sibling "outputs/" would silently resolve to it and
# drop the built files beside the pipeline scripts. Build next to the source.
OUT = os.path.normpath(os.path.join(HERE, "..", "Outputs", "reports"))
os.makedirs(OUT, exist_ok=True)

BODY = 8.8          # tight enough to hold the note to two pages, still readable
TABLE = 8.0
LINE = 1.08

# The theme's rich_run handles **bold** and `code`. The source note also uses
# *italic*, which would otherwise print its own asterisks - so tokenise here
# rather than stripping the markers and losing the emphasis.
_TOK = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*`]+?\*)")


def rich(p, text, size=BODY):
    for tok in _TOK.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.font.name, r.font.bold = T.FONT, True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = T.MONO
            r.font.color.rgb = T.D_BLUE
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.font.name, r.font.italic = T.FONT, True
        else:
            r = p.add_run(tok); r.font.name = T.FONT
        r.font.size = DPt(size)
    return p


def squeeze(doc, pt=2):
    """Shrink the spacer paragraph the table helper appends, to reclaim space."""
    par = doc.paragraphs[-1]
    if par.text.strip():
        return
    par.paragraph_format.space_after = DPt(0)
    par.paragraph_format.space_before = DPt(0)
    (par.runs[0] if par.runs else par.add_run("")).font.size = DPt(pt)


def bullet(doc, text, size=BODY):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.left_indent = docx.shared.Inches(0.25)
    par.paragraph_format.space_after = DPt(3)
    par.paragraph_format.line_spacing = LINE
    return rich(par, text, size=size)


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = DPt(6 if level == 1 else 5)
    p.paragraph_format.space_after = DPt(2)
    return p


def para(doc, text, size=BODY, italic=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = DPt(space_after)
    p.paragraph_format.line_spacing = LINE
    rich(p, text, size=size)
    if italic:
        for r in p.runs:
            r.font.italic = True
    return p


def mono_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = DPt(6)
    p.paragraph_format.left_indent = docx.shared.Inches(0.18)
    r = p.add_run(text)
    r.font.name, r.font.size = T.MONO, DPt(9)
    r.font.color.rgb = T.D_BLUE
    return p


def build() -> str:
    doc = docx.Document()
    T.docx_setup(doc, margin=0.7)

    T.cover(
        doc,
        "Population-weighted access to adequately staffed facilities",
        "Methodological note for peer review",
        meta_lines=[
            "National Health Facility Assessment  ·  620 wards  ·  22,936,947 population",
            "Source of record: Outputs/reports/06_methodological_note.md",
            "Long-form results: 04_access_method_and_results.md  ·  Sensitivity grid: 04_sensitivity_analysis.csv",
        ],
        kicker="Everything needed to reproduce the measure or attack it is in this note",
    )

    # ---------------------------------------------------------------- 1
    h(doc, "1  The estimand", 1)
    T.callout(
        doc, "The measure",
        "The share of ward population living within 60 minutes' one-way travel time of a "
        "facility that meets the published minimum staffing standard for its own facility type.",
    )
    squeeze(doc)
    para(doc,
         "Two things it is deliberately **not**: a measure of distance to *any* facility "
         "(reported alongside, 43.6%, because the difference is the policy point), and a measure "
         "of whether care is actually obtained. It is modelled physical access to a qualifying "
         "supply point.")

    # ---------------------------------------------------------------- 2
    h(doc, "2  Inputs (as received, all synthetic)", 1)
    T.add_table(
        doc,
        ["Input", "Used for", "n"],
        [
            ["health_facilities.csv", "Supply points", "1,346 → 1,315 located; 31 unusable coordinates"],
            ["facility_personnel_scores.mif/.mid", "Staffing by cadre", "544 adequate / 648 inadequate / 123 unassessed"],
            ["minimum_staffing_norms.csv", "Adequacy rule", "5 facility types"],
            ["admin_boundaries.gpkg → wards", "Demand geography and population denominator", "620 wards; 22,936,947 people"],
            ["road_network.geojson", "Network routing", "213 segments → 334 nodes / 426 edges, 1 component"],
        ],
        widths=[30, 30, 40], size=TABLE,
    )
    squeeze(doc)
    para(doc,
         "Population is taken from the boundary layer, not `ward_population.csv`, because the CSV "
         "is missing 14 values and the two agree wherever both exist.")

    # ---------------------------------------------------------------- 3
    h(doc, "3  Method and parameters", 1)
    for i, (lead, rest) in enumerate([
        ("Adequacy",
         " — a facility is adequate if it meets or exceeds the published minimum for **every cadre "
         "with a non-zero minimum for its type**. Zero-minimum cadres are not tested, so a health "
         "post is not failed for lacking a medical officer. No cut point was invented."),
        ("Demand points",
         " — each ward is sampled on a regular grid clipped to its polygon, step scaled to ward area "
         "(`TARGET_SAMPLES_PER_WARD=30`, clamped to 700–12,000 m): **18,924 points, median 30/ward, "
         "range 27–34**. Ward centroids were rejected: they answer “is the middle of this ward served”."),
        ("Travel time",
         " — hybrid, as in AccessMod and the WHO–UNICEF accessibility guidance: "
         "`t = min(walk-to-road + on-network route + walk-from-road, direct off-road)`. Off-road "
         "**5 km/h**. The `min` means a road that detours is correctly ignored."),
        ("Routing",
         " — one multi-source Dijkstra seeded at all facilities (symmetric on an undirected graph, so "
         "identical to per-pair routing). Junctions are snapped to a 1 m grid before use as node "
         "identities, so float noise cannot split a junction and silently disconnect the network. "
         "Analysis CRS **ESRI:102022**, equal-area, distances in metres."),
        ("Coverage",
         " — the *fraction* of a ward's points within threshold, so coverage is continuous. "
         "Population covered = fraction × ward population."),
        ("Second measure",
         " — 2SFCA: each adequate facility's staff ÷ population in its own catchment, summed over the "
         "facilities reaching each ward. Reported as staff per 10,000."),
    ], start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = DPt(2)
        p.paragraph_format.line_spacing = LINE
        p.paragraph_format.left_indent = docx.shared.Inches(0.25)
        p.paragraph_format.first_line_indent = docx.shared.Inches(-0.25)
        r = p.add_run(f"{i}.  ")
        r.font.name, r.font.size, r.font.bold = T.FONT, DPt(BODY), True
        r.font.color.rgb = T.D_BLUE
        r2 = p.add_run(lead)
        r2.font.name, r2.font.size, r2.font.bold = T.FONT, DPt(BODY), True
        rich(p, rest, size=BODY)

    para(doc,
         "**Every parameter is a named constant.** `common.py:108` threshold, `:117` off-road speed, "
         "`:92` CRS, `:125` snap limit, `:111`/`:120` sweep grids; "
         "`04_access_analysis_pipeline.py:85-86` sampling density.",
         space_after=2)

    # ---------------------------------------------------------------- 4
    h(doc, "4  Reproduction", 1)
    mono_block(doc, "python Outputs/run_all.py        # stages 01-05; stage 04 is this measure")
    para(doc,
         "Deterministic: the demand grid is `np.arange`, not sampled — **there is no RNG anywhere, "
         "so there is no seed to set**. Verified on geopandas 1.1.3, networkx 3.6.1, shapely 2.1.2, "
         "scipy 1.16.3, pandas 2.3.3, numpy 2.3.5. Outputs to check: "
         "`data/ward_access_metrics.csv` (620 rows), `data/ward_access_metrics.params.json` (the "
         "parameters that produced it), and `logs/04_access_analysis.log`, which records every "
         "intermediate count quoted above.")

    # ---------------------------------------------------------------- 5
    h(doc, "5  Result, and how far it moves", 1)
    para(doc,
         "**21.6% covered (4,962,704 of 22,936,947); 131 wards at zero.** Full grid in "
         "`04_sensitivity_analysis.csv` — 26 rows, each also split by population-estimation method.")
    T.add_table(
        doc,
        ["Lever", "Range tested", "Covered share"],
        [
            ["Threshold", "30 → 120 min", "5.8% → 62.2%"],
            ["Off-road speed", "4 → 15 km/h", "14.4% → 86.0%"],
            ["Unassessed facilities", "inadequate → adequate", "21.6% → 25.6%"],
            ["Population method", "projection vs gridded", "20.9% vs 22.7%"],
        ],
        widths=[34, 33, 33], size=TABLE,
    )
    squeeze(doc)
    para(doc,
         "**The level is not robust. The ranking is.** Wards at zero coverage at 60 min are "
         "overwhelmingly the same wards at zero at 90 and 120, because their problem is absent "
         "supply, not marginal travel time. Quote the priority list with more confidence than the "
         "percentage.")

    # ---------------------------------------------------------------- 6
    h(doc, "6  How to challenge this", 1)
    bullet(doc,
                  "**The road network barely contributes at the headline threshold.** It supplies the "
                  "fastest route for **0.3%** of covered demand points at 60 min — 0.0% at 30 and 45, "
                  "4.6% at 120. At 5 km/h a 60-minute budget buys about 5 km, and on a 213-segment "
                  "national skeleton the walk to the road usually exceeds that. **The honest "
                  "description of the headline is “population within a ~5 km walk of an adequately "
                  "staffed facility”.** The network binds only at longer thresholds. This is the "
                  "first thing a reviewer should press on, and it is why off-road speed — not the "
                  "threshold — is the most consequential assumption.")
    bullet(doc,
                  "**Reproduce the attack cheaply:** change `OFFROAD_SPEED_KMH` (`common.py:117`) and "
                  "rerun stage 04, about 30 seconds. Every headline number in the report is "
                  "regenerated from the sweep, not typed, so nothing goes stale.")
    bullet(doc,
                  "**Directional bias is knowable.** Unmapped feeder roads mean the model under-uses "
                  "roads that really exist, so **21.6% is a lower bound**.")

    # ---------------------------------------------------------------- 7
    h(doc, "7  Known weaknesses, ranked by effect", 1)
    for lead, rest in [
        ("Uniform population within wards",
         " — the largest error. No gridded population is supplied. Understates coverage where people "
         "cluster near roads or facilities, overstates it where they cluster away. Unremovable with "
         "these data; only nameable."),
        ("Off-road speed is doing most of the work",
         " (§6). A 4 → 15 km/h change moves the headline 72 points."),
        ("The road layer is a skeleton",
         " — 213 segments for a country. On-network times are pessimistic."),
        ("Speeds are static",
         " — no seasonality, rainy-season impassability, congestion, ferries or river crossings. In "
         "this setting seasonal impassability is likely the largest missing term."),
        ("123 unassessed and 31 unlocated facilities",
         " — both bounded rather than hidden, but the 31 unlocated are excluded entirely and *cannot* "
         "be bounded. If they cluster in poorly served wards, coverage is understated by an unknown "
         "amount."),
        ("Supply-side only",
         " — assumes the facility is open, stocked and functional on arrival, and that a staffed post "
         "is a substitute for a hospital. It is not."),
        ("2SFCA inherits every assumption above",
         ", plus the premise that catchment population is a fair divisor of staffed capacity."),
    ]:
        par = doc.add_paragraph(style="List Number")
        par.paragraph_format.left_indent = docx.shared.Inches(0.25)
        par.paragraph_format.space_after = DPt(2)
        par.paragraph_format.line_spacing = LINE
        rl = par.add_run(lead)
        rl.font.name, rl.font.size, rl.font.bold = T.FONT, DPt(BODY), True
        rich(par, rest, size=BODY)

    # ---------------------------------------------------------------- 8
    h(doc, "8  Errata in 04_access_method_and_results.md", 1)
    para(doc,
         "Found while preparing this note. Both are reporting defects, not analysis defects, and "
         "neither changes a result.", space_after=4)
    bullet(doc,
                  "**§7 states the 90th:10th percentile 2SFCA ratio as “9500934990.9×”.** The 10th "
                  "percentile is **exactly 0.00** — 131 wards have no accessible adequate staff — so "
                  "the ratio is undefined and the printed figure is a division-by-zero artefact. "
                  "Correct statement: *the 90th-percentile ward has 9.50 staff per 10,000; the "
                  "10th-percentile ward has none, so the ratio is unbounded.* Fix the generator to "
                  "test for a zero denominator.")
    bullet(doc,
                  "**§4 prose says “124 unknown facilities”; its own table and the stage-04 log say "
                  "123.** 123 is correct.")

    path = os.path.join(OUT, "06_methodological_note.docx")
    doc.save(path)
    return path


if __name__ == "__main__":
    p = build()
    print("wrote", p, f"({os.path.getsize(p) / 1024:.0f} KB)")
